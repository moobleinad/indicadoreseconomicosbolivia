import os
import docx

md_path = r'c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz\0_Plan_Estrategico_Posicionamiento_y_Contenidos_DanielSimons.md'
docx_path = r'c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz\0_Plan_Estrategico_Posicionamiento_y_Contenidos_DanielSimons.docx'

content_md = """# PLAN ESTRATÉGICO DE POSICIONAMIENTO, MODELO DE NEGOCIO Y CONTENIDOS POR EMBUDO
**DANIEL SIMONS**
*De ideas complejas a resultados concretos.*
**IDEAS • ESTRUCTURA • IMPACTO**
*Documento Maestro Consolidado (01 de Agosto de 2026)*

---

## 1. PROPÓSITO & DECLARACIÓN DE MARCA
Ayudar a personas, jóvenes y organizaciones a transformar ideas complejas en proyectos, modelos, estrategias y conocimiento que generen resultados concretos.

> **«No soy una agencia de marketing. No soy un coach. No soy únicamente consultor. Soy un estructurador de ideas complejas: construyo claridad.»**

---

## 2. PÚBLICO OBJETIVO Y 4 SEGMENTOS DE MERCADO

### 📌 Segmento 1: Emprendedores y Dueños de MYPE (Sector Productivo)
* **Perfil:** Dueños de micro y pequeñas empresas en Santa Cruz y Bolivia.
* **Dolor:** Negocios desordenados, falta de estructura en costos y modelos de negocio.
* **Solución:** FORJA (De la idea al proyecto) e Impulso MYPE.

### 📌 Segmento 2: Profesionales Independientes y Ejecutivos (Sector Conocimiento)
* **Perfil:** Consultores, directores, investigadores, médicos, abogados.
* **Dolor:** Gran conocimiento acumulado pero disperso sin lograr plasmarlo en libros o informes técnicos.
* **Solución:** DESTILADO (Del conocimiento disperso a la claridad).

### 🎓 Segmento 3: Jóvenes Estudiantes (BTH / Universidad) y su Ecosistema (Padres, Directores, Docentes)
* **Perfil:**
  1. *Estudiantes:* Colegios con Bachillerato Técnico Humanístico (BTH), institutos y universitarios preparando proyectos de grado o ferias.
  2. *Padres de Familia:* Buscan formación real en emprendimiento para sus hijos.
  3. *Directores & Docentes:* Buscan metodologías probadas para ferias productivas BTH y proyectos de grado.
* **Dolor:** Proyectos escolares improvisados, pánico a la tesis, falta de guías prácticas.
* **Solución:** EL JUEGO DEL EMPRENDEDOR, GUÍA SOBREVIVIENDO A LA TESIS, Talleres BTH.

### 🏛️ Segmento 4: Instituciones, Gremios y Empresas (Sector Estratégico)
* **Perfil:** Cámaras empresariales, gremios, instituciones educativas y grupos de decisión.
* **Dolor:** Abundancia de opiniones políticas y falta de datos económicos verificados.
* **Solución:** PROYECTOS PROPIOS (Observatorio Económico Empresarial & MFEIR).

---

## 3. MATRIZ DE CONTENIDOS POR EMBUDO DE CONVERSIÓN

```
   EMBUDO DE CONVERSIÓN DE AUTORIDAD (DANIEL SIMONS)
   
  ┌──────────────────────────────────────────────────────────┐
  │ TOFU: Atracción & Posicionamiento (Redes / Blog)         │
  │ - Cuestionar la falta de estructura en MYPEs.             │
  │ - Criticar la improvisación en proyectos BTH/Escolares.  │
  │ - Análisis estructural de la economía boliviana.        │
  └──────────────────────────┬───────────────────────────────┘
                             │
  ┌──────────────────────────▼───────────────────────────────┐
  │ MOFU: Demostración de Método & Educación (Artículos Blog) │
  │ - Explicar CÓMO funciona el Método de 6 semanas BTH.     │
  │ - Enseñar el proceso del Destilado y la Forja.           │
  │ - Metodología práctica para sobrevivir a la Tesis.        │
  └──────────────────────────┬───────────────────────────────┘
                             │
  ┌──────────────────────────▼───────────────────────────────┐
  │ BOFU: Cierre & Solicitud de Servicio / Compra de Libro    │
  │ - Venta directa de "El Juego del Emprendedor".           │
  │ - Captación de clientes para Forja y Destilado de Ideas.  │
  └──────────────────────────────────────────────────────────┘
```

---

## 4. CALENDARIO EDITORIAL DE REDACCIÓN (FASE 1)
1. **Artículo 1 (BTH & Colegios):** *«El drama del Bachillerato Técnico Humanístico (BTH): Por qué los proyectos escolares terminan en la basura y cómo estructurar emprendimientos reales.»*
2. **Artículo 2 (MYPEs):** *«Por qué publicar en TikTok no salvará a tu MYPE: El error de buscar marketing cuando te falta estructura de negocio.»*
3. **Artículo 3 (Universitarios):** *«Sobreviviendo a la Tesis sin morir en el intento: La guía práctica para estructurar tu proyecto de grado sin pánico.»*
4. **Artículo 4 (Ejecutivos):** *«De 100 páginas de información dispersa a un documento técnico ejecutivo de 5 páginas (El proceso del Destilado).»*
"""

# Save Markdown
with open(md_path, 'w', encoding='utf-8') as f:
    f.write(content_md)

# Save Word DOCX
doc = docx.Document()
doc.add_heading("PLAN ESTRATÉGICO DE POSICIONAMIENTO Y CONTENIDOS - DANIEL SIMONS", 0)
doc.add_paragraph("De ideas complejas a resultados concretos. IDEAS • ESTRUCTURA • IMPACTO")
doc.add_paragraph("-" * 50)

for line in content_md.split('\n'):
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.strip():
        doc.add_paragraph(line.strip())

doc.save(docx_path)
print("SUCCESS: Master strategic plan saved in both MD and DOCX formats!")
