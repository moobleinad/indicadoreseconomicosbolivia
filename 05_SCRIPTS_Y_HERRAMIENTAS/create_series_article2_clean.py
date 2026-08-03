import os
import docx
from docx.shared import Pt, RGBColor, Inches

root_dir = r"c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz"
art_base_dir = os.path.join(root_dir, "02_ARTICULOS_Y_PUBLICACIONES")
dir_art_serie2 = os.path.join(art_base_dir, "02.04_Articulo_Serie_2_Flotacion_Libre_vs_Flotacion_Sucia")
os.makedirs(dir_art_serie2, exist_ok=True)

md_path = os.path.join(dir_art_serie2, "02.04_Articulo_Serie_2_Flotacion_Libre_vs_Flotacion_Sucia.md")
docx_path = os.path.join(dir_art_serie2, "02.04_Articulo_Serie_2_Flotacion_Libre_vs_Flotacion_Sucia.docx")

# CLEAN MD CONTENT (SHORT PARAGRAPHS, SYSTEMIC FLOW, ZERO EMOJIS, ZERO '---' LINES)
content_md = """# RÉGIMEN CAMBIARIO: EL DEBATE ENTRE LA TERAPIA DE CHOQUE Y LA FLOTACIÓN GRADUAL
### Serie Política Monetaria y Sociedad | Parte 2
### Por Daniel Simons

El paso repentino hacia un tipo de cambio totalmente flexible ha generado una sacudida inmediata en la economía boliviana. Pasar sin anestesia de una cotización congelada durante más de una década a una flotación libre transfiere toda la volatilidad cambiaria directamente al bolsillo del ciudadano y a la estructura de costos de las pequeñas y medianas empresas.

Al igual que ocurrió con el incremento brusco en el precio de los combustibles, este salto repentino obliga a la economía privada y a las familias a pagar de golpe los grandes desajustes acumulados por el Estado.

1. La cadena de transmisión del choque libre

En una economía endeble con baja disponibilidad de divisas en ventanilla bancaria, la flotación libre desencadena un mecanismo de transmisión inmediato.

La incertidumbre sobre el valor diario del dólar eleva los precios de reposición de insumos e importaciones. Este incremento se traslada en tiempo real a los bienes de consumo final, acelerando la presión inflacionaria.

El resultado es la erosión del poder adquisitivo y el agotamiento del capital de trabajo de las MYPEs, que ven cómo su liquidez se evapora al intentar reponer inventario.

2. La alternativa técnica: El modelo de flotación sucia administrada

Frente a una terapia de choque, la teoría monetaria y la experiencia de la región demuestran que existía un camino de ajuste predecible: el modelo de flotación sucia o deslizamiento progresivo con bandas cambiarias.

Bajo este esquema, el Banco Central establece una pauta de ajuste gradual conocido de antemano por el mercado.

Este mecanismo permite que importadores, productores y familias proyecten sus costos de reposición y flujo de caja con margen de planificación, evitando picos especulativos salvajes y situaciones de pánico.

3. Adaptación estructural: Cómo proteger tu patrimonio y tu empresa

Con el régimen flexible en marcha, la solución para las empresas y familias no es esperar medidas gubernamentales milagrosas, sino adaptar la estructura interna.

Resguardar excedentes líquidos en divisas duras o en activos reales tangibles preserva el valor del capital frente al deterioro de la moneda nacional.

Para las empresas, revisar el margen bruto real, ajustar los ciclos de cobro y adaptar la propuesta de valor a la nueva realidad de costos es la única garantía de continuidad.

Si necesitas evaluar la estructura financiera de tu empresa, proteger tu liquidez o rediseñar tu modelo operativo frente al nuevo escenario cambiario, conoce nuestro programa de acompañamiento técnico:

[EVALUAR ESTRUCTURA DE MI NEGOCIO CON DANIEL SIMONS](https://www.danielsimons.xyz/p/impulso-mype-360.html)
"""

with open(md_path, "w", encoding="utf-8") as f:
    f.write(content_md)

# Save DOCX
doc = docx.Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p_title = doc.add_paragraph()
run_title = p_title.add_run("RÉGIMEN CAMBIARIO: EL DEBATE ENTRE LA TERAPIA DE CHOQUE Y LA FLOTACIÓN GRADUAL")
run_title.font.name = "Montserrat"
run_title.font.size = Pt(16)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(188, 167, 114)

p_sub = doc.add_paragraph()
run_sub = p_sub.add_run("Serie Política Monetaria y Sociedad | Parte 2\nPor Daniel Simons")
run_sub.font.name = "Montserrat"
run_sub.font.size = Pt(11)
run_sub.font.bold = True
run_sub.font.color.rgb = RGBColor(128, 128, 128)

for line in content_md.split("\n"):
    if line.startswith("# "):
        continue
    elif line.strip():
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Montserrat"
        r.font.size = Pt(10.5)

doc.save(docx_path)
print("SUCCESS: Series Article 2 Clean MD and DOCX Created!")
