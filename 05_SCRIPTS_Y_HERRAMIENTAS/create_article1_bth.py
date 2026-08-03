import os
import docx

md_art1_path = r'c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz\0_Articulo_1_El_Drama_del_BTH_DanielSimons.md'
docx_art1_path = r'c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz\0_Articulo_1_El_Drama_del_BTH_DanielSimons.docx'

art1_text = """# El drama del Bachillerato Técnico Humanístico (BTH): Por qué los proyectos escolares terminan en la basura y cómo estructurar emprendimientos reales

**Por Daniel Simons**  
*Estructurador de Ideas Complejas | Docente de Emprendimiento e Historia CIFA*

---

Cada año en Santa Cruz y en toda Bolivia, miles de estudiantes de secundaria se enfrentan al reto del **Bachillerato Técnico Humanístico (BTH)** y a las ferias escolares de emprendimiento. Durante semanas, padres de familia, docentes y alumnos invierten tiempo, dinero en maquetas, impresiones y compras de insumos para presentar un "proyecto productivo".

Sin embargo, cuando la feria termina y los diplomas se entregan, ocurre una triste realidad: **el 95% de esos proyectos terminan en la basura o archivados en un cajón**.

¿Por qué sucede esto si en Bolivia sobra la creatividad y las ganas de salir adelante? 

La respuesta no es la falta de talento de los jóvenes. El problema es que el sistema educativo les enseña a **improvisar una manualidad o un producto**, pero nunca les enseña la **estructura real de un modelo de negocio**.

---

## 1. Los 3 Errores Fatales del BTH Tradicional

### Error 1: Confundir una "Manualidad" con un Proyecto Productivo
En la mayoría de los colegios se le pide al estudiante que "fabrique algo". El alumno hace galletas, mermeladas, jabones o manualidades con reciclaje. Se enfoca el 100% del esfuerzo en la elaboración física del objeto, pero cero por ciento en responder las preguntas cruciales: ¿A quién le resuelve un problema real? ¿Cuánto cuesta producirlo exactamente? ¿Cómo se va a comercializar fuera de las paredes del colegio?

### Error 2: La Falsa Motivación sin Estructura de Costos
Muchos docentes intentan entusiasmar a los estudiantes con discursos motivacionales de "sé tu propio jefe". Pero cuando el estudiante compra los ingredientes en el mercado, olvida calcular el costo de su propio tiempo, la energía eléctrica, el empaque y el transporte. El resultado es un producto que se vende a pérdida o que no es competitivo en el mercado real de Santa Cruz.

### Error 3: El Pánico a la Exposición y la Falta de Argumentación
Llegado el día de la evaluación, los estudiantes memorizan un discurso tipo recitación. Si el jurado les hace una pregunta fuera del guión sobre el precio o la competencia, se paralizan. Esto ocurre porque memorizaron texto en lugar de **comprender la estructura lógica de su proyecto**.

---

## 2. La Solución Estructural: El Método de 6 Semanas (De la Idea al Impacto Real)

Como docente e investigador en la enseñanza de emprendimientos productivos, he desarrollado y comprobado una metodología estructurada de 6 pasos para que un estudiante de secundario o instituto pase de una idea dispersa a un proyecto real:

1. **Semana 1 - El Diagnóstico del Problema:** No empieces por el producto; empieza por identificar una necesidad no resuelta en tu barrio o ciudad.
2. **Semana 2 - La Propuesta de Valor Única:** Define en una sola frase por qué alguien te compraría a ti y no a la competencia.
3. **Semana 3 - La Estructura de Costos Rigurosa:** Aprende a calcular el costo unitario real, el margen de ganancia y el precio final justo.
4. **Semana 4 - El Modelo de Negocio MFEIR Simplificado:** Diseña cómo se conecta tu oferta con tu cliente final sin depender de favores familiares.
5. **Semana 5 - La Validación en Mercado Real:** Pon a prueba tu prototipo con 10 personas ajenas a tu familia antes de gastar en la feria.
6. **Semana 6 - La Presentación de Alto Impacto:** Aprende a defender tu proyecto con seguridad, datos y argumentos sólidos ante cualquier jurado.

---

## 3. Un Mensaje para los Padres de Familia y Directores

A los **Padres de Familia**: El objetivo de un proyecto BTH no es que gastes dinero a última hora comprando materiales para que tu hijo cumpla por una nota. El objetivo es que tu hijo adquiera **criterio económico, disciplina y estructura mental** para la vida adulta.

A los **Directores y Catedráticos**: Cuando le damos a los jóvenes las herramientas correctas de estructuración, dejamos de evaluar carpetas aburridas y empezamos a ver nacer verdaderos emprendedores y líderes.

---

## 💡 ¿Quieres estructurar proyectos reales en tu colegio o universidad?

No le des a tus hijos o estudiantes discursos motivacionales vacíos: **dales estructura y claridad**.

* 📘 **Adquiere el libro «El Juego del Emprendedor»:** La guía práctica diseñada especialmente para jóvenes y estudiantes.
* 🎓 **Talleres & Capacitaciones BTH para Colegios:** Capacitaciones estructuradas para docentes y alumnos.

📲 **Contacto directo con Daniel Simons:** [WhatsApp de Consultoría y Formación](https://wa.me/59170000000)  
🌐 **Portal Oficial:** [www.danielsimons.xyz](https://www.danielsimons.xyz)
"""

with open(md_art1_path, 'w', encoding='utf-8') as f:
    f.write(art1_text)

# Save Word DOCX
doc = docx.Document()
doc.add_heading("El drama del Bachillerato Técnico Humanístico (BTH)", 0)
doc.add_paragraph("Por Daniel Simons - Estructurador de Ideas Complejas")
doc.add_paragraph("-" * 50)

for line in art1_text.split('\n'):
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[3:], level=3)
    elif line.strip():
        doc.add_paragraph(line.strip())

doc.save(docx_art1_path)
print("SUCCESS: Article 1 (BTH) saved in MD and DOCX formats!")
