import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_whatsapp_templates():
    folder_path = r"c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz\10_PLAN_1_PROMOCION_INDICADORES"
    os.makedirs(folder_path, exist_ok=True)

    md_path = os.path.join(folder_path, "10.02_Plantillas_Mensajes_WhatsApp.md")
    docx_path = os.path.join(folder_path, "10.02_Plantillas_Mensajes_WhatsApp.docx")

    dashboard_url = "https://www.danielsimons.xyz/p/indicadores-economicos-de-bolivia_0349188327.html"
    channel_url = "https://whatsapp.com/channel/0029VbDAeCQ1t90gu0qjtC07"

    md_content = f"""# PLANTILLAS DE MENSAJES OFICIALES PARA EL CANAL DE WHATSAPP

**Propietario:** Lic. Daniel Simons  
**Sitio Web:** [www.danielsimons.xyz](https://www.danielsimons.xyz/)  
**Página de Indicadores:** [{dashboard_url}]({dashboard_url})  
**Canal Oficial de WhatsApp:** [{channel_url}]({channel_url})  

---

## 📢 PLANTILLA 1: MENSAJE CORTO DE ANUNCIO OFICIAL DE HORARIOS

*Usar para anunciar a los seguidores la programación fija de las 2 actualizaciones diarias.*

```text
🇧🇴 INDICADORES ECONÓMICOS DE BOLIVIA
📢 Horarios Oficiales de Actualización Diaria

Estimada comunidad ejecutiva y empresarial, este canal monitorea los datos clave de Bolivia en 2 horarios estratégicos:

🌅 08:45 AM | Resumen Matutino (Dólar P2P / Mercado Libre)
Captura la apertura y primera cotización consolidada de insumos e importaciones.

🌙 20:00 PM | Reporte Nocturno (BCB / Cierre Oficial)
Registra la cotización oficial del Banco Central tras el cierre de operaciones interbancarias.

📊 Consulta el tablero interactivo en vivo los 365 días:
👉 {dashboard_url}

Compilado por Daniel Simons | www.danielsimons.xyz
```

---

## 📊 PLANTILLA 2: RESUMEN EJECUTIVO DIARIO DE INDICADORES

*Usar para la difusión diaria con las cifras actualizadas.*

```text
🇧🇴 INDICADORES ECONÓMICOS DE BOLIVIA
📅 Resumen Ejecutivo • 03/08/2026
Compilado por Daniel Simons | www.danielsimons.xyz

🟢 Dólar Libre (P2P): 11.75 Bs (Estabilizado)
🔴 Dólar Oficial (BCB): 12.13 Bs (Ajuste Flexible)
🔴 Inflación IPC: 4.82% 1S (Interanual: 9.23%)
🟡 Reservas RIN: $3.617,3 MM (Oro: $2.882,9M)
🟢 Balanza Comercial: +$1.669 MM (Superávit)
🟡 Riesgo País (EMBI): 430 pbs (Moderado)
🟢 Oro Exportación: $2.450 USD/oz (Histórico Alto)
🟡 Crecimiento PIB: 2.10% (Moderado)

💡 Reflexión Coyuntural del Día:
"Con la convergencia del dólar libre a 11.75 Bs, ¿tu empresa ya recalculó el precio real de reposición de inventarios?"

📊 Revisa el tablero interactivo completo:
👉 {dashboard_url}

📲 Comparte esta información con otros ejecutivos.
```

---
"""

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)

    print(f"Archivo Markdown de plantillas creado en: {md_path}")

    # Generación de archivo Word (.docx)
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("PLANTILLAS DE MENSAJES PARA CANAL DE WHATSAPP")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(188, 167, 114)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run(f"Propietario: Lic. Daniel Simons | Sitio Web: www.danielsimons.xyz\nCanal: {channel_url}")
    run_meta.font.name = "Arial"
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True
    run_meta.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    h_sec1 = doc.add_heading(level=1)
    run_hsec1 = h_sec1.add_run("📢 PLANTILLA 1: MENSAJE CORTO DE ANUNCIO OFICIAL DE HORARIOS")
    run_hsec1.font.name = "Arial"
    run_hsec1.font.color.rgb = RGBColor(188, 167, 114)

    anuncio_text = (
        "🇧🇴 INDICADORES ECONÓMICOS DE BOLIVIA\n"
        "📢 Horarios Oficiales de Actualización Diaria\n\n"
        "Estimada comunidad ejecutiva y empresarial, este canal monitorea los datos clave de Bolivia en 2 horarios estratégicos:\n\n"
        "🌅 08:45 AM | Resumen Matutino (Dólar P2P / Mercado Libre)\n"
        "Captura la apertura y primera cotización consolidada de insumos e importaciones.\n\n"
        "🌙 20:00 PM | Reporte Nocturno (BCB / Cierre Oficial)\n"
        "Registra la cotización oficial del Banco Central tras el cierre de operaciones interbancarias.\n\n"
        f"📊 Consulta el tablero interactivo en vivo los 365 días:\n👉 {dashboard_url}\n\n"
        "Compilado por Daniel Simons | www.danielsimons.xyz"
    )

    p_anuncio = doc.add_paragraph(anuncio_text)
    p_anuncio.runs[0].font.name = "Arial"
    p_anuncio.runs[0].font.size = Pt(10.5)

    h_sec2 = doc.add_heading(level=1)
    run_hsec2 = h_sec2.add_run("📊 PLANTILLA 2: RESUMEN EJECUTIVO DIARIO DE INDICADORES")
    run_hsec2.font.name = "Arial"
    run_hsec2.font.color.rgb = RGBColor(188, 167, 114)

    diario_text = (
        "🇧🇴 INDICADORES ECONÓMICOS DE BOLIVIA\n"
        "📅 Resumen Ejecutivo • 03/08/2026\n"
        "Compilado por Daniel Simons | www.danielsimons.xyz\n\n"
        "🟢 Dólar Libre (P2P): 11.75 Bs (Estabilizado)\n"
        "🔴 Dólar Oficial (BCB): 12.13 Bs (Ajuste Flexible)\n"
        "🔴 Inflación IPC: 4.82% 1S (Interanual: 9.23%)\n"
        "🟡 Reservas RIN: $3.617,3 MM (Oro: $2.882,9M)\n"
        "🟢 Balanza Comercial: +$1.669 MM (Superávit)\n"
        "🟡 Riesgo País (EMBI): 430 pbs (Moderado)\n"
        "🟢 Oro Exportación: $2.450 USD/oz (Histórico Alto)\n"
        "🟡 Crecimiento PIB: 2.10% (Moderado)\n\n"
        "💡 Reflexión Coyuntural del Día:\n"
        "\"Con la convergencia del dólar libre a 11.75 Bs, ¿tu empresa ya recalculó el precio real de reposición de inventarios?\"\n\n"
        f"📊 Revisa el tablero interactivo completo:\n👉 {dashboard_url}\n\n"
        "📲 Comparte esta información con otros ejecutivos."
    )

    p_diario = doc.add_paragraph(diario_text)
    p_diario.runs[0].font.name = "Arial"
    p_diario.runs[0].font.size = Pt(10.5)

    doc.save(docx_path)
    print(f"Archivo Word de plantillas creado en: {docx_path}")

if __name__ == "__main__":
    create_whatsapp_templates()
