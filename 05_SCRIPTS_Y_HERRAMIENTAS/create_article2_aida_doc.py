import os
import docx
from docx.shared import Pt, RGBColor, Inches

root_dir = r"c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz"
art_dir = os.path.join(root_dir, "02_ARTICULOS_Y_PUBLICACIONES")

md_path = os.path.join(art_dir, "02.02_Articulo_2_Por_Que_TikTok_No_Salvara_Tu_MYPE.md")
docx_path = os.path.join(art_dir, "02.02_Articulo_2_Por_Que_TikTok_No_Salvara_Tu_MYPE.docx")

content_md = """# POR QUÉ PUBLICAR EN TIKTOK NO SALVARÁ A TU MYPE: EL ERROR DE BUSCAR MARKETING CUANDO TE FALTA ESTRUCTURA DE NEGOCIO
### Por Daniel Simons | Estructurador de Ideas Complejas

---

En los últimos meses, conversaciones con propietarios de pequeñas y medianas empresas (MYPEs) y emprendedores en Bolivia revelan un patrón recurrente: frente a la presión de la inflación, el aumento de costos de reposición y la brecha del dólar paralelo, la respuesta inmediata suele ser la misma: *«necesito más marketing, necesito hacer más videos en TikTok para vender más»*.

Sin embargo, en el contexto económico actual, buscar más ventas sin haber ordenado primero la estructura financiera y operativa del negocio no solo es ineficaz; es peligroso.

---

### 📌 1. ATENCIÓN: La ilusión de las ventas frente al dólar paralelo

Cuando los insumos importados, la materia prima y los repuestos se adquieren indexados a la cotización del dólar en el mercado paralelo (~12 Bs.), pero los precios de venta al público se cobran en moneda nacional sin un análisis riguroso de margen real, vender más no genera mayor liquidez.

Vender más bajo una estructura descalibrada únicamente acelera el agotamiento de capital de trabajo. Muchas MYPEs están vendiendo con entusiasmo su stock actual, para luego descubrir que el dinero recaudado no les alcanza para comprar la mitad de la mercadería de reposición.

---

### 💡 2. INTERÉS: Marketing vs. Estructura de Negocio

El marketing es un amplificador, no un corrector de deficiencias estructurales. Si el motor de un vehículo está fallando por falta de lubricación o desajuste mecánico, presionar el acelerador al máximo solo provocará que el motor se funda más rápido.

Publicar contenidos virales en redes sociales puede traer clientes a la puerta, pero si tu empresa carece de:

1. **Cálculo preciso del Costo de Reposición:** Basado en el valor real de reposición de insumos, no en el costo histórico de compra.
2. **Control de Flujo de Caja y Liquidez:** Separando claramente el capital operativo de los fondos de reserva.
3. **Estandarización de Procesos Internos:** Eliminando desperdicios y fugas invisibles de dinero en la operación diaria.

Entonces cada nuevo cliente atraído por el marketing estará consumiendo un margen que en realidad no existe.

---

### 🎯 3. DESEO: Transformar el desorden en claridad y sostenibilidad

La verdadera competitividad de una MYPE en tiempos de incertidumbre cambiaria no proviene de adivinar tendencias en redes, sino de **construir una estructura de negocio sólida y resiliente**.

Ordenar tu empresa requiere tres pasos fundamentales:

- **Saber dónde estás parado (Diagnóstico Real):** Evaluar el margen bruto real ajustado por inflación e insumos importados.
- **Proteger la caja antes de buscar expansión:** Priorizar la liquidez y la rotación eficiente de inventario sobre el crecimiento apresurado.
- **Rediseñar la propuesta de valor:** Competir por criterio, claridad y eficiencia operativa, en lugar de competir por precio bajo en un mercado encarecido.

---

### 🚀 4. ACCIÓN: Construye la estructura que tu negocio necesita

Antes de invertir tiempo y recursos en campañas agresivas de publicidad, asegúrate de que la estructura sobre la que descansa tu empresa sea capaz de sostener ese crecimiento.

El orden no es un lujo para las grandes corporaciones; es el salvavidas fundamental de la pequeña y mediana empresa.

*Si necesitas evaluar la estructura de tu negocio, ordenar tus costos de reposición o rediseñar tu modelo operativo para proteger tu liquidez, conoce más sobre nuestro acompañamiento técnico a MYPEs:*

👉 **[CONOCER EL SERVICIO DE IMPULSO MYPE 360°](https://www.danielsimons.xyz/p/impulso-mype-360.html)**
"""

with open(md_path, "w", encoding="utf-8") as f:
    f.write(content_md)

# Create DOCX
doc = docx.Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p_title = doc.add_paragraph()
run_title = p_title.add_run("POR QUÉ PUBLICAR EN TIKTOK NO SALVARÁ A TU MYPE: EL ERROR DE BUSCAR MARKETING CUANDO TE FALTA ESTRUCTURA DE NEGOCIO")
run_title.font.name = "Montserrat"
run_title.font.size = Pt(16)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(188, 167, 114)

p_sub = doc.add_paragraph()
run_sub = p_sub.add_run("Por Daniel Simons | Estructurador de Ideas Complejas")
run_sub.font.name = "Montserrat"
run_sub.font.size = Pt(11)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph("="*60)

for line in content_md.split("\n"):
    if line.startswith("### "):
        h = doc.add_heading(line.replace("### ", ""), level=3)
        for r in h.runs:
            r.font.name = "Montserrat"
            r.font.color.rgb = RGBColor(188, 167, 114)
    elif line.startswith("---") or line.startswith("# "):
        continue
    elif line.strip():
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Montserrat"
        r.font.size = Pt(10.5)

doc.save(docx_path)
print("SUCCESS: Article 2 AIDA docx and md created successfully!")
