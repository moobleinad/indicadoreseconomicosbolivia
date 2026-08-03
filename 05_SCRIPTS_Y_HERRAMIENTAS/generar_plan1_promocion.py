import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def update_plan1_files():
    folder_path = r"c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz\10_PLAN_1_PROMOCION_INDICADORES"
    os.makedirs(folder_path, exist_ok=True)

    md_path = os.path.join(folder_path, "10.01_Plan_1_Promocion_Indicadores_Economicos.md")
    docx_path = os.path.join(folder_path, "10.01_Plan_1_Promocion_Indicadores_Economicos.docx")

    dashboard_url = "https://www.danielsimons.xyz/p/indicadores-economicos-de-bolivia_0349188327.html"
    channel_url = "https://whatsapp.com/channel/0029VbDAeCQ1t90gu0qjtC07"

    md_content = f"""# PLAN 1: ESTRATEGIA Y PROMOCIÓN DE LOS INDICADORES ECONÓMICOS DE BOLIVIA

**Propietario:** Lic. Daniel Simons  
**Sitio Web:** [www.danielsimons.xyz](https://www.danielsimons.xyz/)  
**Página de Indicadores:** [{dashboard_url}]({dashboard_url})  
**Canal Oficial de WhatsApp:** [{channel_url}]({channel_url})  
**Eslogan Principal de Captación:** ACTUALIZACIONES DIARIAS: UNIRSE AL CANAL DE WHATSAPP  
**Versión del Tema Activa:** v31_theme_optimizado.xml  
**Estado:** EN EJECUCIÓN ACTIVA  

---

## 🎯 OBJETIVO PRINCIPAL
Convertir el Dashboard de Indicadores Económicos de Bolivia y el Canal Oficial de WhatsApp en la fuente de referencia macroeconómica diaria indispensable para empresarios, ejecutivos, consultores e inversores en Bolivia, ofreciendo datos actualizados, análisis de causa-efecto y reflexiones estratégicas antes del inicio de la jornada laboral (07:00 AM).

---

## 🚀 SECCIÓN 1: CANALES Y PASOS DE EJECUCIÓN DEL PLAN 1

### 1. 🌐 Integración en el Sitio Web (`www.danielsimons.xyz`)
- **Portada Principal (Tema v31):** Banner oficial con título a la izquierda, botón `[ Ver Indicadores ➔ ]` arriba a la derecha, y botón verde `[ 📲 ACTUALIZACIONES DIARIAS: UNIRSE AL CANAL DE WHATSAPP ➔ ]` abarcando la fila inferior centrada.
- **Pie del Dashboard & Header:** Botón discreto y elegante: `[ 📲 ACTUALIZACIONES DIARIAS: UNIRSE AL CANAL DE WHATSAPP ➔ ]` enlazando directamente a `{channel_url}`.
- **Artículos y Ensayos:** Llamado a la acción (CTA) al pie de cada artículo económico: *"ACTUALIZACIONES DIARIAS: UNIRSE AL CANAL DE WHATSAPP para recibir el resumen de las 07:00 AM."*

### 2. 💼 Estrategia Orgánica en LinkedIn (Perfil Profesional de Daniel Simons)
- **Frecuencia:** 2 a 3 publicaciones semanales enfocadas en coyuntura macroeconómica y gestión de proyectos.
- **Contenido:** Reflexión técnica sobre datos recientes (Dólar P2P a 11.75 Bs, Inflación acumulada 4.82%, Riesgo País a 430 pbs).
- **Gancho:** *"¿Cómo impacta la convergencia del tipo de cambio libre en la estructura de costos operacionales de tu empresa?"*
- **Llamado a la Acción (CTA en 1er comentario):** `📲 ACTUALIZACIONES DIARIAS: UNIRSE AL CANAL DE WHATSAPP ({channel_url})`.

### 3. 📱 Estados de WhatsApp y Difusión en Grupos Empresariales
- **Estados de WhatsApp:** Publicación diaria de la tarjeta visual de indicadores con la fecha cronológica del día en vivo y el sticker de enlace directo al Canal `{channel_url}` bajo el texto: *"ACTUALIZACIONES DIARIAS: UNIRSE AL CANAL DE WHATSAPP"*.
- **Grupos Seleccionados:** Difusión respetuosa en grupos de graduados, cámaras empresariales (CAINCO, FEPSC, etc.) y asociaciones profesionales.

### 4. 🔍 Optimización de Búsqueda Orgánica en WhatsApp (Novedades)
- **Nombre Oficial del Canal:** `Indicadores Económicos Bolivia | Daniel Simons`
- **Enlace de Invitación:** `{channel_url}`
- **Palabras Clave en Descripción:** Dólar Bolivia, Tipo de Cambio Libre, Inflación, BCB, Balanza Comercial, Reservas RIN, Economía Bolivia.
- **Posicionamiento:** Facilita que cualquier usuario en Bolivia encuentre el canal mediante el buscador interno de WhatsApp.

### 5. 📧 Firma Institucional de Correo & Propuestas Comerciales
- Inclusión del enlace del canal `{channel_url}` bajo la leyenda *"ACTUALIZACIONES DIARIAS: UNIRSE AL CANAL DE WHATSAPP"* en la firma de correo electrónico de Daniel Simons y en propuestas de consultoría.

---

## 📜 SECCIÓN 2: REGLAS ESPECÍFICAS Y DIRECTIVAS DEL LIC. DANIEL SIMONS

1. **Regla de Maquetación del Banner de Portada (Tema v31):** El banner oficial de portada mantiene el título a la izquierda, el botón `[ Ver Indicadores ➔ ]` en la parte superior derecha, y el botón verde de WhatsApp `[ 📲 ACTUALIZACIONES DIARIAS: UNIRSE AL CANAL DE WHATSAPP ➔ ]` en una fila inferior propia, abarcando el ancho de forma totalmente centrada.
2. **Regla de Eslogan / Leyenda Oficial de Captación:** El texto y llamado a la acción oficial para todos los botones, banners, firmas y pies de publicaciones es estrictamente: **"ACTUALIZACIONES DIARIAS: UNIRSE AL CANAL DE WHATSAPP"**.
3. **Regla de Enlace Oficial del Canal:** El enlace único y oficial del Canal de WhatsApp para toda integración en web, publicaciones, redes y firmas es: `{channel_url}`.
4. **Regla de Estructura de Tarjetas (Regla #9):** Las tarjetas de cada indicador se mantienen 100% limpias, técnicas e institucionales. Las preguntas reflexivas se ubican exclusivamente en los espacios libres ENTRE cada cuadro.
5. **Regla de Fecha Cronológica en Vivo (Regla #10):** Cada tarjeta debe mostrar explícitamente la fecha cronológica del día actual en vivo ("Actualizado al día: [Fecha Actual]"), independiente de la fecha del período de referencia oficial del indicador.
6. **Regla de Ejecución del Plan:** Cada vez que el Lic. Daniel Simons indique "veamos", "revisemos" o "ejecutemos el plan 1", se consultará y presentará de inmediato este documento con sus actualizaciones.

---
"""

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)

    print(f"Archivo Markdown actualizado en: {md_path}")

    # Generar Word
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("PLAN 1: ESTRATEGIA Y PROMOCIÓN DE LOS INDICADORES ECONÓMICOS DE BOLIVIA")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(188, 167, 114)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run(f"Propietario: Lic. Daniel Simons | Sitio Web: www.danielsimons.xyz\nCanal WhatsApp: {channel_url}\nTema Activo: v31_theme_optimizado.xml\nEstado: EN EJECUCIÓN ACTIVA")
    run_meta.font.name = "Arial"
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True
    run_meta.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    h_obj = doc.add_heading(level=1)
    run_hobj = h_obj.add_run("🎯 OBJETIVO PRINCIPAL")
    run_hobj.font.name = "Arial"
    run_hobj.font.color.rgb = RGBColor(188, 167, 114)

    p_obj = doc.add_paragraph(
        "Convertir el Dashboard de Indicadores Económicos de Bolivia y el Canal Oficial de WhatsApp en la fuente de referencia macroeconómica diaria indispensable para empresarios, ejecutivos, consultores e inversores en Bolivia, ofreciendo datos actualizados, análisis de causa-efecto y reflexiones estratégicas antes del inicio de la jornada laboral (07:00 AM)."
    )
    p_obj.runs[0].font.name = "Arial"
    p_obj.runs[0].font.size = Pt(11)

    h_sec1 = doc.add_heading(level=1)
    run_hsec1 = h_sec1.add_run("🚀 SECCIÓN 1: CANALES Y PASOS DE EJECUCIÓN DEL PLAN 1")
    run_hsec1.font.name = "Arial"
    run_hsec1.font.color.rgb = RGBColor(188, 167, 114)

    pasos = [
        ("1. 🌐 Integración en el Sitio Web (www.danielsimons.xyz)",
         f"• Portada Principal (Tema v31): Banner oficial con título a la izquierda, botón Ver Indicadores arriba a la derecha, y botón verde de WhatsApp en la fila inferior centrada.\n"
         f"• Pie del Dashboard & Header: Botón discreto [ 📲 ACTUALIZACIONES DIARIAS: UNIRSE AL CANAL DE WHATSAPP ➔ ] enlazando a {channel_url}.\n"
         "• Artículos y Ensayos: Llamado a la acción (CTA) al pie de cada artículo económico."),

        ("2. 💼 Estrategia Orgánica en LinkedIn (Perfil Profesional)",
         f"• Frecuencia: 2 a 3 publicaciones semanales sobre coyuntura macroeconómica.\n"
         "• Contenido: Reflexión técnica sobre datos recientes (Dólar P2P, Inflación, Riesgo País).\n"
         "• Gancho: Reflexiones sobre estructura de costos operacionales.\n"
         f"• CTA en 1er comentario: ACTUALIZACIONES DIARIAS: UNIRSE AL CANAL DE WHATSAPP ({channel_url})."),

        ("3. 📱 Estados de WhatsApp y Difusión en Grupos Empresariales",
         f"• Estados de WhatsApp: Publicación diaria de la tarjeta visual con leyenda \"ACTUALIZACIONES DIARIAS: UNIRSE AL CANAL DE WHATSAPP\" ({channel_url}).\n"
         "• Grupos Seleccionados: Difusión respetuosa en grupos de graduados, cámaras empresariales y colegios profesionales."),

        ("4. 🔍 Optimización de Búsqueda Orgánica en WhatsApp (Novedades)",
         f"• Nombre Oficial: Indicadores Económicos Bolivia | Daniel Simons\n"
         f"• Enlace de Invitación: {channel_url}\n"
         "• Palabras Clave: Dólar Bolivia, Tipo de Cambio Libre, Inflación, BCB, Balanza Comercial."),

        ("5. 📧 Firma Institucional de Correo & Propuestas Comercial",
         f"• Inclusión de \"ACTUALIZACIONES DIARIAS: UNIRSE AL CANAL DE WHATSAPP\" ({channel_url}) en la firma de correo electrónico y contraportada de propuestas.")
    ]

    for titulo, desc in pasos:
        h_item = doc.add_heading(level=2)
        run_hitem = h_item.add_run(titulo)
        run_hitem.font.name = "Arial"
        run_hitem.font.size = Pt(13)
        run_hitem.font.color.rgb = RGBColor(40, 40, 40)

        p_desc = doc.add_paragraph(desc)
        p_desc.runs[0].font.name = "Arial"
        p_desc.runs[0].font.size = Pt(10.5)

    h_sec2 = doc.add_heading(level=1)
    run_hsec2 = h_sec2.add_run("📜 SECCIÓN 2: REGLAS ESPECÍFICAS Y DIRECTIVAS DEL LIC. DANIEL SIMONS")
    run_hsec2.font.name = "Arial"
    run_hsec2.font.color.rgb = RGBColor(188, 167, 114)

    reglas = [
        "Regla de Maquetación del Banner de Portada (Tema v31): El banner oficial de portada mantiene el título a la izquierda, el botón [ Ver Indicadores ➔ ] en la parte superior derecha, y el botón verde de WhatsApp [ 📲 ACTUALIZACIONES DIARIAS: UNIRSE AL CANAL DE WHATSAPP ➔ ] en una fila inferior propia, abarcando el ancho de forma totalmente centrada.",
        "Regla de Eslogan / Leyenda Oficial de Captación: El texto y llamado a la acción oficial para todos los botones, banners, firmas y pies de publicaciones es strictly: \"ACTUALIZACIONES DIARIAS: UNIRSE AL CANAL DE WHATSAPP\".",
        f"Regla de Enlace Oficial del Canal: El enlace único y oficial del Canal de WhatsApp para toda integración en web, publicaciones, redes y firmas es: {channel_url}.",
        "Regla de Estructura de Tarjetas (Regla #9): Las tarjetas de cada indicador se mantienen 100% limpias, técnicas e institucionales. Las preguntas reflexivas se ubican exclusivamente en los espacios libres ENTRE cada cuadro.",
        "Regla de Fecha Cronológica en Vivo (Regla #10): Cada tarjeta debe mostrar explícitamente la fecha cronológica del día actual en vivo (\"Actualizado al día: [Fecha Actual]\"), independiente de la fecha del período de referencia oficial.",
        "Regla de Ejecución del Plan: Cada vez que el Lic. Daniel Simons indique \"veamos\", \"revisemos\" o \"ejecutemos el plan 1\", se consultará y presentará de inmediato este documento con sus actualizaciones."
    ]

    for r in reglas:
        p_r = doc.add_paragraph(r, style='List Bullet')
        p_r.runs[0].font.name = "Arial"
        p_r.runs[0].font.size = Pt(10.5)

    doc.save(docx_path)
    print(f"Archivo Word actualizado en: {docx_path}")

if __name__ == "__main__":
    update_plan1_files()
