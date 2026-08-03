import os
import docx
from docx.shared import Pt, RGBColor, Inches

root_dir = r"c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz"
master_md_path = os.path.join(root_dir, "00_DOCUMENTOS_ESTRATÉGICOS", "00.05_Prompt_Master_Diseno_y_Reglas_Web_DanielSimons.md")
master_docx_path = os.path.join(root_dir, "00_DOCUMENTOS_ESTRATÉGICOS", "00.05_Prompt_Master_Diseno_y_Reglas_Web_DanielSimons.docx")

content_md = """# 📜 PROMPT MASTER & REGLAS INMUTABLES DE LA WEB
### Portal Oficial: [www.danielsimons.xyz](https://www.danielsimons.xyz/)

---

### 1. 🎯 POSICIONAMIENTO DE MARCA
- **Marca:** Daniel Simons — Estructurador de Ideas Complejas (IDEAS • ESTRUCTURA • IMPACTO).
- **Estilo Visual:** Ejecutivo, neutral, sobrio y elegante. Sin retórica política, sin gráficos infantiles y sin iconos estilo años 80.

---

### 2. ⚡ AUTORIZACIÓN TOTAL DE EJECUCIÓN DE SCRIPTS Y PROTOCOLO "DALE"
- **EJECUCIÓN DIRECTA SIN PEDIR PERMISOS NUNCA:** Prohibido interrumpir con solicitudes emergentes o pedir permisos para ejecutar comandos de consola o scripts de Python. El asistente tiene **autorización absoluta y total** para ejecutar Python, comandos internos y actualizar archivos de forma fluida e ininterrumpida.
- **PROTOCOLO DE CONFIRMACIÓN CON "DALE":** Para avanzar a la redacción, modificación o publicación de contenido estratégico, el asistente presentará la propuesta e ideas clave y **esperará siempre que Daniel le diga "DALE" o le dé su visto bueno explícito** antes de proceder.

---

### 3. ✍️ GUÍA INMUTABLE DE ESTILO EDITORIAL DANIEL SIMONS
- **PÁRRAFOS CORTOS Y REDACCIÓN SISTÉMICA:**
  * **Párrafos Cortos:** Máximo 2 a 4 líneas por bloque. Cero bloques masivos de texto. Escaneo visual ultra-rápido en móviles y PCs.
  * **Redacción Sistémica:** Articular las ideas como un sistema interconectado de causa y efecto (Norma → Cambio Monetario → Liquidez Bancaria → Costo de Reposición → Ajuste Empresarial/Social).
- **CERO VERBORREA Y CERO HISTORIA ABURRIDA:** Prohibido prólogos académicos, introducciones históricas largas o discursos patronizantes. Entrada directa y quirúrgica desde la primera línea.
- **PROHIBICIÓN ABSOLUTA DE ICONOS, EMOJIS Y LÍNEAS SEPARADORAS (`---`):** NUNCA JAMÁS colocar emojis ni iconos en títulos o texto. NUNCA JAMÁS colocar rayas separadoras (`---` / `<hr/>`) dentro del cuerpo de los artículos.
- **TONO DE ALTA AUTORIDAD Y CERO CRÍTICA PATRONIZANTE:** Trato de igual a igual con el lector. Respeto absoluto al emprendedor y ciudadano. Enfoque en dar criterios de acción, no juzgar.
- **CONVERSIÓN UTILITARIA ORGÁNICA:** Cada escrito conduce de forma natural hacia el siguiente artículo o servicio de estructuración (*WhatsApp / Impulso MYPE / Destilado de Ideas*).

---

### 4. 📷 ESTÁNDAR FOTOGRÁFICO DEFINITIVO Y FORMATO HORIZONTAL (16:9 MULTIREDES)
- **FORMATO OFICIAL HORIZONTAL (16:9):** Las imágenes oficiales de artículos y páginas se generan en formato **Horizontal (16:9)** para permitir su uso óptimo en Blogger, OpenGraph de Facebook, LinkedIn, Twitter/X y previews de WhatsApp sin cortes indeseados.
- **PROTOCOLO DE SUBIDA MANUAL EN PÁGINAS ESTÁTICAS (`/p/*.html`):**
  * Para cualquier página estática o panel (Indicadores Económicos, Servicios, Quienes Somos, etc.), el asistente genera la imagen oficial Horizontal 16:9 en `.webp`.
  * **Daniel sube la imagen manualmente al final de la página dentro del editor de Blogger** para garantizar que Blogger la registre nativamente como la miniatura oficial de previsualización en WhatsApp y redes sociales.
  * **Daniel comparte el enlace CDN oficial (`blogger.googleusercontent.com`)** para confirmación.
- **ADAPTACIÓN A PORTADA (MINIATURA CUADRADA 1:1):** Para el carrusel de la portada principal de la web, la misma imagen horizontal se adapta quirúrgicamente a 1:1 (`object-fit: cover`).
- **ENFOQUE DE CONTENIDO (ESTILO ALT2):** Macro-realismo utilitario enfocado en la materia del artículo (manos en primer plano manipulando divisas en dólares y bolivianos, cuadernos de cuero negro con anotaciones estratégicas, herramientas financieras directas, cero estereotipos).
- **ILUMINACIÓN Y COLORES (ESTILO ALT5 EQUIPETROL ATARDECER):**
  - **Iluminación:** Luz crepuscular de atardecer en tonos dorados cálidos reflejándose sobre superficies oscuras, vidrio y cuero negro.
  - **Paleta y Sombras:** Sombras profundas en negro absoluto (`#000000`) combinadas con destellos y reflejos dorados calientes (`#BCA772`).
- **Tipografía Móvil:** Títulos de tarjetas en celulares fijados en **`15px !important` (Negrita)**.

---

### 5. 📐 ESTRUCTURA DE LA PORTADA (6 SECCIONES EN CARRUSEL HORIZONTAL)
La portada se compone de filas con deslizamiento horizontal en carrusel (flechas doradas `❮` `❯`):

1. **ÚLTIMAS ENTRADAS DEL BLOG:** Carrusel de artículos con botón *"Ver todas las entradas ➔"*.
2. **SERVICIOS DE ESTRUCTURACIÓN:** *Destilado de Ideas*, *FORJA*, *Marketing 360°*, *Impulso MYPE*.
3. **PROYECTOS PROPIOS & INVESTIGACIÓN:** *Modelo MFEIR*, *El Juego del Emprendedor*.
4. **LIBROS:** Sección dedicada a publicaciones y obras de Daniel Simons.
5. **QUIÉN SOY (DANIEL SIMONS):** Tarjeta ejecutiva de presentación.
6. **TRABAJOS Y PROPUESTAS:** *Propuesta Urbanizaciones*, *Transparencia Inteligente*, *Guía Sobreviviendo a la Tesis*, *Evaluar Desarrollo de Marca*.

---

### 6. 📁 ESTRUCTURA DE ARCHIVOS EN CARPETAS DEDICADAS Y SISTEMA NUMÉRICO
- **CARPETAS INDIVIDUALES POR ARTÍCULO / PÁGINA:** Cada artículo o proyecto se organizará dentro de su propia carpeta dedicada conteniendo sus archivos `.md`, `.docx` e imagen `.webp` oficial **Horizontal (16:9)**.
- **Jerarquía Numérica:** Prefijos numéricos secuenciales (`00_...`, `01_...`, `02_...`).

---

### 7. ⚡ LÍMITES TÉCNICOS Y PESO EN BLOGGER
- **Peso Máximo del XML:** El archivo XML del tema NUNCA DEBE SUPERAR LOS 200 KB (Rango ideal: 160 KB – 190 KB).
- **CLÁUSULA DE CONTROL ABSOLUTO:** Sólo modificarás LO QUE EXPRESAMENTE SE TE INSTRUYA; TODO LO DEMÁS QUEDARÁ EXACTAMENTE IGUAL.
"""

with open(master_md_path, "w", encoding="utf-8") as f:
    f.write(content_md)

# Save Master DOCX
doc = docx.Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p_title = doc.add_paragraph()
run_title = p_title.add_run("PROMPT MASTER & REGLAS INMUTABLES DE LA WEB")
run_title.font.name = "Montserrat"
run_title.font.size = Pt(20)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(188, 167, 114)

p_sub = doc.add_paragraph()
run_sub = p_sub.add_run("Portal Oficial: www.danielsimons.xyz | Daniel Simons")
run_sub.font.name = "Montserrat"
run_sub.font.size = Pt(11)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph("="*60)

def add_sec(title, text_lines):
    h = doc.add_heading(title, level=2)
    for r in h.runs:
        r.font.name = "Montserrat"
        r.font.color.rgb = RGBColor(188, 167, 114)
    for line in text_lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Montserrat"
        r.font.size = Pt(10.5)

add_sec("1. POSICIONAMIENTO DE MARCA", [
    "• Marca: Daniel Simons — Estructurador de Ideas Complejas (IDEAS • ESTRUCTURA • IMPACTO).",
    "• Estilo Visual: Ejecutivo, neutral, sobrio y elegante. Sin retórica política, sin gráficos infantiles y sin iconos estilo años 80."
])

add_sec("2. AUTORIZACIÓN TOTAL DE EJECUCIÓN DE SCRIPTS Y PROTOCOLO 'DALE'", [
    "• EJECUCIÓN DIRECTA SIN PEDIR PERMISOS NUNCA: Prohibido interrumpir con solicitudes emergentes o pedir permisos para ejecutar comandos de consola o scripts de Python.",
    "• PROTOCOLO DE CONFIRMACIÓN CON 'DALE': Para avanzar a la redacción, modificación o publicación de contenido estratégico, el asistente esperará siempre que Daniel le diga 'DALE'."
])

add_sec("3. GUÍA INMUTABLE DE ESTILO EDITORIAL DANIEL SIMONS", [
    "• PÁRRAFOS CORTOS: Máximo 2 a 4 líneas por bloque.",
    "• REDACCIÓN SISTÉMICA: Articular ideas como un sistema interconectado de causa y efecto.",
    "• CERO VERBORREA Y CERO EMOJIS/LÍNEAS: Prohibición absoluta de emojis e iconos."
])

add_sec("4. ESTÁNDAR FOTOGRÁFICO Y PROTOCOLO DE SUBIDA EN PÁGINAS ESTÁTICAS (/p/*.html)", [
    "• FORMATO OFICIAL HORIZONTAL (16:9): Las imágenes oficiales de artículos y páginas se generan en formato Horizontal (16:9).",
    "• PROTOCOLO DE SUBIDA MANUAL EN PÁGINAS ESTÁTICAS: Para cualquier página estática o panel (Indicadores Económicos, Servicios, etc.), el asistente genera la imagen oficial Horizontal 16:9 en .webp.",
    "• Daniel sube la imagen manualmente al final de la página dentro del editor de Blogger para que Blogger la registre nativamente como la miniatura oficial de previsualización en WhatsApp y redes sociales.",
    "• Daniel comparte el enlace CDN oficial (blogger.googleusercontent.com) para confirmación."
])

add_sec("5. ESTRUCTURA DE LA PORTADA (CARRUSEL HORIZONTAL)", [
    "1. ÚLTIMAS ENTRADAS DEL BLOG: Carrusel de artículos con botón 'Ver todas las entradas ➔'.",
    "2. SERVICIOS DE ESTRUCTURACIÓN: Destilado de Ideas, FORJA, Marketing 360°, Impulso MYPE.",
    "3. PROYECTOS PROPIOS & INVESTIGACIÓN: Modelo MFEIR, El Juego del Emprendedor.",
    "4. LIBROS: Sección oficial de libros y publicaciones.",
    "5. QUIÉN SOY (DANIEL SIMONS): Tarjeta ejecutiva de presentación.",
    "6. TRABAJOS Y PROPUESTAS: Urbanizaciones, Transparencia Inteligente, Guía Tesis, Evaluar Marca."
])

add_sec("6. ESTRUCTURA DE ARCHIVOS EN CARPETAS DEDICADAS Y SISTEMA NUMÉRICO", [
    "• CARPETAS INDIVIDUALES POR ARTÍCULO / PÁGINA: Cada artículo o proyecto se organizará dentro de su propia carpeta dedicada conteniendo sus archivos .md, .docx e imagen .webp oficial Horizontal (16:9)."
])

add_sec("7. LÍMITES TÉCNICOS Y PESO EN BLOGGER", [
    "• Peso Máximo del XML: NUNCA DEBE SUPERAR LOS 200 KB (Rango ideal: 160 KB – 190 KB).",
    "• CLÁUSULA DE CONTROL ABSOLUTO: Sólo modificarás LO QUE EXPRESAMENTE SE TE INSTRUYA; TODO LO DEMÁS QUEDARÁ EXACTAMENTE IGUAL."
])

doc.save(master_docx_path)
print("SUCCESS: Updated Master MD and DOCX with Static Page Image Upload Protocol!")
