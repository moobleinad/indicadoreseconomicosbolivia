import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_plan2_files():
    folder_path = r"c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz\11_PLAN_2_AUTOMATIZACION_NUBE"
    os.makedirs(folder_path, exist_ok=True)

    md_path = os.path.join(folder_path, "11.01_Plan_2_Automatizacion_en_la_Nube_Blogger_y_WhatsApp.md")
    docx_path = os.path.join(folder_path, "11.01_Plan_2_Automatizacion_en_la_Nube_Blogger_y_WhatsApp.docx")

    dashboard_url = "https://www.danielsimons.xyz/p/indicadores-economicos-de-bolivia_0349188327.html"
    channel_url = "https://whatsapp.com/channel/0029VbDAeCQ1t90gu0qjtC07"

    md_content = f"""# PLAN 2: AUTOMATIZACIÓN TOTAL EN LA NUBE 24/7 (BLOGGER Y WHATSAPP)

**Propietario:** Lic. Daniel Simons  
**Sitio Web:** [www.danielsimons.xyz](https://www.danielsimons.xyz/)  
**Página de Indicadores:** [{dashboard_url}]({dashboard_url})  
**Canal Oficial de WhatsApp:** [{channel_url}]({channel_url})  
**Estado:** REGISTRADO Y LISTO PARA EJECUCIÓN  

---

## 🎯 OBJETIVO PRINCIPAL DEL PLAN 2
Garantizar la actualización macroeconómica diaria 24/7 en la nube tanto para el sitio web `www.danielsimons.xyz` como para el Canal Oficial de WhatsApp, sin depender en absoluto de que la computadora personal del Lic. Daniel Simons esté encendida.

---

## 🚀 SECCIÓN 1: LOS 2 CAMBIOS CLAVE DEL PLAN 2

### 1. ☁️ CAMBIO 1: Automatización en la Nube de Blogger (GitHub Actions / Cloud Python 24/7)
- **Independencia de hardware:** Migrar la ejecución del script `auto_update_indicadores.py` a un servidor en la nube 100% gratuito (GitHub Actions / Google Cloud Cron).
- **Horarios Programados Nativos:**  
  - 🌅 **08:45 AM BOT:** Verificación matutina de la apertura del mercado dólar P2P y fecha del día en vivo.  
  - 🌙 **20:00 PM BOT:** Verificación nocturna del reporte oficial de cotizaciones del Banco Central de Bolivia (BCB) y cierre de oro en LBMA.
- **Resultado:** La web de Blogger se actualizará sola de forma ininterrumpida los 365 días del año.

### 2. 📲 CAMBIO 2: Automatización e Integración Directa con el Canal de WhatsApp (Meta Cloud API)
- **Conexión de API Oficial:** Configurar la conexión con la API Oficial de Meta (WhatsApp Cloud API) de forma 100% gratuita (dentro del límite sin costo de 1.000 conversaciones/mes de Meta).
- **Publicación Matutina (07:00 AM):** El servidor en la nube enviará automáticamente el resumen ejecutivo formateado directamente al Canal de WhatsApp (`{channel_url}`) antes de la apertura del mercado.
- **Formato del Mensaje:**  
  - Encabezado institucional: `🇧🇴 INDICADORES ECONÓMICOS DE BOLIVIA`  
  - Fecha del día en vivo.  
  - Los 8 indicadores con sus semáforos RAG (`🟢`, `🟡`, `🔴`).  
  - La pregunta de reflexión coyuntural del día.  
  - Enlace al Dashboard interactivo completo.

---

## 📜 SECCIÓN 2: REGLAS ESPECÍFICAS Y DIRECTIVAS DEL LIC. DANIEL SIMONS

1. **Regla de Cero Dependencia de PC:** La arquitectura del Plan 2 debe operar de forma 100% autónoma en servidores de la nube, sin requerir intervención manual ni que la computadora personal esté encendida.
2. **Regla de Costo Cero:** La implementación utilizará exclusivamente las capas gratuitas oficiales de GitHub Actions / Google Cloud y Meta Cloud API.
3. **Regla de Ejecución del Plan 2:** Cada vez que el Lic. Daniel Simons indique "veamos", "revisemos" o "ejecutemos el plan 2", se abrirá y desplegará de inmediato este documento con los pasos detallados de configuración.

---
"""

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)

    print(f"Archivo Markdown del Plan 2 creado en: {md_path}")

    # Generación de archivo Word (.docx)
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("PLAN 2: AUTOMATIZACIÓN TOTAL EN LA NUBE 24/7 (BLOGGER Y WHATSAPP)")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(188, 167, 114)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run(f"Propietario: Lic. Daniel Simons | Sitio Web: www.danielsimons.xyz\nCanal WhatsApp: {channel_url}\nEstado: REGISTRADO Y LISTO PARA EJECUCIÓN")
    run_meta.font.name = "Arial"
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True
    run_meta.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    h_obj = doc.add_heading(level=1)
    run_hobj = h_obj.add_run("🎯 OBJETIVO PRINCIPAL DEL PLAN 2")
    run_hobj.font.name = "Arial"
    run_hobj.font.color.rgb = RGBColor(188, 167, 114)

    p_obj = doc.add_paragraph(
        "Garantizar la actualización macroeconómica diaria 24/7 en la nube tanto para el sitio web www.danielsimons.xyz como para el Canal Oficial de WhatsApp, sin depender en absoluto de que la computadora personal del Lic. Daniel Simons esté encendida."
    )
    p_obj.runs[0].font.name = "Arial"
    p_obj.runs[0].font.size = Pt(11)

    h_sec1 = doc.add_heading(level=1)
    run_hsec1 = h_sec1.add_run("🚀 SECCIÓN 1: LOS 2 CAMBIOS CLAVE DEL PLAN 2")
    run_hsec1.font.name = "Arial"
    run_hsec1.font.color.rgb = RGBColor(188, 167, 114)

    pasos = [
        ("1. ☁️ CAMBIO 1: Automatización en la Nube de Blogger (GitHub Actions / Cloud Python 24/7)",
         "• Independencia de hardware: Migrar la ejecución del script auto_update_indicadores.py a un servidor en la nube 100% gratuito.\n"
         "• Horarios Programados Nativos: 08:45 AM BOT (Matutino P2P) y 20:00 PM BOT (Nocturno BCB).\n"
         "• Resultado: La web se actualizará sola de forma ininterrumpida los 365 días del año."),

        ("2. 📲 CAMBIO 2: Automatización e Integración Directa con el Canal de WhatsApp (Meta Cloud API)",
         f"• Conexión de API Oficial: Configurar Meta WhatsApp Cloud API de forma 100% gratuita.\n"
         f"• Publicación Matutina (07:00 AM): Envío automático del resumen ejecutivo directo al Canal de WhatsApp ({channel_url}).\n"
         "• Formato: Encabezado institucional, fecha del día en vivo, los 8 indicadores con semáforos RAG y reflexión del día.")
    ]

    for titulo, desc in pasos:
        h_item = doc.add_heading(level=2)
        run_hitem = h_item.add_run(titulo)
        run_hitem.font.name = "Arial"
        run_hitem.font.size = Pt(13)
        run_hitem.font.color.rgb = RGBColor(40, 40, 40)

        p_desc = doc.add_paragraph(desc)
        p_desc.runs[0].font.name = "Arial"
        p_desc.runs[0].font.size = Pt(10.5)

    h_sec2 = doc.add_heading(level=1)
    run_hsec2 = h_sec2.add_run("📜 SECCIÓN 2: REGLAS ESPECÍFICAS Y DIRECTIVAS DEL LIC. DANIEL SIMONS")
    run_hsec2.font.name = "Arial"
    run_hsec2.font.color.rgb = RGBColor(188, 167, 114)

    reglas = [
        "Regla de Cero Dependencia de PC: La arquitectura del Plan 2 debe operar de forma 100% autónoma en servidores de la nube, sin requerir intervención manual ni que la computadora personal esté encendida.",
        "Regla de Costo Cero: La implementación utilizará exclusivamente las capas gratuitas oficiales de GitHub Actions / Google Cloud y Meta Cloud API.",
        "Regla de Ejecución del Plan 2: Cada vez que el Lic. Daniel Simons indique \"veamos\", \"revisemos\" o \"ejecutemos el plan 2\", se abrirá y desplegará de inmediato este documento con los pasos detallados de configuración."
    ]

    for r in reglas:
        p_r = doc.add_paragraph(r, style='List Bullet')
        p_r.runs[0].font.name = "Arial"
        p_r.runs[0].font.size = Pt(10.5)

    doc.save(docx_path)
    print(f"Archivo Word del Plan 2 creado en: {docx_path}")

if __name__ == "__main__":
    create_plan2_files()
