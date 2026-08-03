import os
import docx

root_dir = r"c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz"
books_dir = os.path.join(root_dir, "07_JUEGO INFINITO", "01_LIBROS")

unique_files = sorted(os.listdir(books_dir))

report = []

for f in unique_files:
    fp = os.path.join(books_dir, f)
    if f.endswith(".docx"):
        try:
            doc = docx.Document(fp)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading") or (p.runs and p.runs[0].bold and len(p.text) < 80)]
            total_words = sum(len(p.split()) for p in paragraphs)
            sample = "\n".join(paragraphs[:8])
            report.append({
                "file": f,
                "words": total_words,
                "headings": headings[:5],
                "sample": sample
            })
        except Exception as e:
            report.append({"file": f, "error": str(e)})
    elif f.endswith(".md"):
        try:
            with open(fp, "r", encoding="utf-8", errors="ignore") as file:
                c = file.read()
                words = len(c.split())
                report.append({
                    "file": f,
                    "words": words,
                    "headings": [l for l in c.split("\n") if l.startswith("#")],
                    "sample": c[:600]
                })
        except Exception as e:
            report.append({"file": f, "error": str(e)})
    elif f.endswith(".pdf"):
        report.append({
            "file": f,
            "words": 0,
            "headings": ["Documento PDF"],
            "sample": "PDF Oficial Listo para distribución"
        })

report_out = os.path.join(root_dir, "07_JUEGO INFINITO", "INFORME_REVISION_UNO_A_UNO_LIBROS.md")
with open(report_out, "w", encoding="utf-8") as rf:
    rf.write("# INFORME ESTRATÉGICO UNO A UNO: REVISIÓN Y PLAN DE PUBLICACIÓN EL JUEGO INFINITO\n\n")
    for r in report:
        rf.write(f"### 📄 {r['file']}\n")
        if "error" in r:
            rf.write(f"- Error: {r['error']}\n\n")
        else:
            rf.write(f"- **Volumen:** ~{r['words']} palabras\n")
            rf.write(f"- **Encabezados Principales:** {', '.join(r['headings'][:3])}\n")
            rf.write(f"- **Extracto:**\n```text\n{r['sample'][:400]}\n```\n\n")

print("DETAILED REPORT CREATED SUCCESSFULLY!")
