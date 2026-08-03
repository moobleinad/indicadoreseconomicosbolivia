import os
import re
import html
import time
from googleapiclient.discovery import build
from google.oauth2.credentials import Credentials
from googleapiclient.errors import HttpError
import docx

BLOG_ID = '433667097766389126'
TOKEN_FILE = r'c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz\token.json'
SCOPES = ['https://www.googleapis.com/auth/blogger']
BACKUP_DIR = r'c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz\ENTRADAS DESCARTADAS'

os.makedirs(BACKUP_DIR, exist_ok=True)

def get_service():
    creds = Credentials.from_authorized_user_file(TOKEN_FILE, SCOPES)
    return build('blogger', 'v3', credentials=creds)

def clean_filename(text):
    clean = re.sub(r'[\\/*?:"<>|]', "", text)
    return clean.strip()[:100]

def html_to_plain_text(raw_html):
    clean = re.sub(r'<br\s*/?>', '\n', raw_html)
    clean = re.sub(r'</p>', '\n\n', clean)
    clean = re.sub(r'</h[1-6]>', '\n\n', clean)
    clean = re.sub(r'<[^>]+>', '', clean)
    return html.unescape(clean).strip()

def save_to_docx(title, content_html, url):
    filename = f"{clean_filename(title)}.docx"
    filepath = os.path.join(BACKUP_DIR, filename)
    
    doc = docx.Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"URL Original: {url}")
    doc.add_paragraph("-" * 40)
    
    plain_text = html_to_plain_text(content_html)
    for paragraph in plain_text.split('\n\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
            
    doc.save(filepath)
    print(f"  [DOCX SAVED] {filepath}")

def api_call_with_retry(func):
    for attempt in range(5):
        try:
            return func()
        except HttpError as e:
            if e.resp.status == 429:
                wait_time = (attempt + 1) * 3
                print(f"  [QUOTA LIMIT] Rate limit hit. Waiting {wait_time}s before retry...")
                time.sleep(wait_time)
            else:
                raise e
    return func()

def main():
    service = get_service()
    posts = api_call_with_retry(lambda: service.posts().list(blogId=BLOG_ID, maxResults=100).execute()).get('items', [])
    pages = api_call_with_retry(lambda: service.pages().list(blogId=BLOG_ID).execute()).get('items', [])
    
    existing_page_titles = [pg['title'].strip().lower() for pg in pages]
    
    # Categorization Rules
    cat_a_titles = [
        "destilado de ideas de negocio",
        "forja: disciplina, hábitos y lucidez",
        "forja:  disciplina, hábitos y lucidez",
        "el juego del emprendedor: libro para jóvenes",
        "marketing 360",
        "impulso mype 360°",
        "impulso mype 360",
        "liberalismo vs socialismo",
        "propuesta urbanizaciones",
        "transparencia inteligente",
        "propuesta: estrategia integral para la campaña electoral 2026 y análisis de datos en santa cruz, bolivia",
        "evaluar el desarrollo de marca de tu empresa",
        "guía sobreviviendo a la tesis",
        "guia sobreviviendo a la tesis",
        "aportes y comentarios"
    ]

    cat_b_titles = [
        "se abre la votación oficial del concurso fotográfico de la larga noche de museos 2026 cifa",
        "bolivia, que se abra el debate: ¿potencial hidroeléctrico o impacto ambiental?"
    ]

    cat_c_titles = [
        "guía básica: corrientes de pensamiento económico y político en la historia",
        "debate: ¿tiktok destruye el pensamiento crítico o potencia la creatividad y el aprendizaje rápido?",
        "cómo ganar una discusión y defender ideas: el arte de debatir, argumentar y detectar mentiras",
        "diseño inicial (evaluación e tu idea de negocio) superior noche",
        "ideas de negocio formato simplificado",
        "segmentación de mercado para superior noche",
        "ideas de proyectos o emprendimientos (2do borrador) superior tarde",
        "ideas de proyectos o emprendimientos (2do borrador) superior noche",
        "mapas mentales ley 348 feria cifa"
    ]

    print("=== INICIANDO PROCESAMIENTO CON RETRY Y DELAYS ===")

    for p in posts:
        pid = p['id']
        title = p['title'].strip()
        title_lower = title.lower()
        content = p.get('content', '')
        url = p.get('url', '')

        print(f"\nPROCESANDO: '{title}' (ID: {pid})")

        is_cat_a = any(cat in title_lower or title_lower in cat for cat in cat_a_titles)
        is_cat_b = any(cat in title_lower or title_lower in cat for cat in cat_b_titles)
        is_cat_c = any(cat in title_lower or title_lower in cat for cat in cat_c_titles)

        if is_cat_a:
            print("  -> Categoría: PASAR A PÁGINA Y PONER EN BORRADOR")
            if title_lower not in existing_page_titles:
                page_body = {"kind": "blogger#page", "title": title, "content": content}
                new_pg = api_call_with_retry(lambda: service.pages().insert(blogId=BLOG_ID, body=page_body).execute())
                existing_page_titles.append(title_lower)
                print(f"  [PÁGINA CREADA] {new_pg['title']} -> {new_pg.get('url')}")
            else:
                print("  [PÁGINA YA EXISTÍA]")

            api_call_with_retry(lambda: service.posts().revert(blogId=BLOG_ID, postId=pid).execute())
            print("  [ENTRADA CAMBIADA A BORRADOR]")

        elif is_cat_b:
            print("  -> Categoría: PONER EN BORRADOR EN ENTRADAS")
            api_call_with_retry(lambda: service.posts().revert(blogId=BLOG_ID, postId=pid).execute())
            print("  [ENTRADA CAMBIADA A BORRADOR]")

        elif is_cat_c:
            print("  -> Categoría: RESPALDAR EN WORD Y PONER EN BORRADOR")
            save_to_docx(title, content, url)
            api_call_with_retry(lambda: service.posts().revert(blogId=BLOG_ID, postId=pid).execute())
            print("  [ENTRADA CAMBIADA A BORRADOR]")

        else:
            print("  -> No categorizado específicamente. Se mantiene intacto.")

        time.sleep(1.5)

    print("\n=== PROCESAMIENTO COMPLETADO CON ÉXITO ===")

if __name__ == "__main__":
    main()
