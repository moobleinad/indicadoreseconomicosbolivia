import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles & Fonts
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    # Colors
    GOLD = RGBColor(0x99, 0x7A, 0x15) # Warm Gold / Bronze
    DARK_BLUE = RGBColor(0x1B, 0x36, 0x5D)
    GRAY = RGBColor(0x55, 0x55, 0x55)
    
    # Title Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("INFORME DE AUDITORÍA Y EVALUACIÓN WEB\nDANIELSIMONS.XYZ")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = DARK_BLUE
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Análisis Técnico, UI/UX, SEO, Estrategia de IA e Integración con Blogger")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = GOLD
    
    # Meta table / Box
    tbl_meta = doc.add_table(rows=1, cols=1)
    tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl_meta.cell(0, 0)
    set_cell_background(cell, "F4F6F9")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p_m = cell.paragraphs[0]
    p_m.paragraph_format.space_after = Pt(2)
    p_m.add_run("Sitio Web Auditado: ").bold = True
    p_m.add_run("https://www.danielsimons.xyz\n")
    p_m.add_run("Fecha de Evaluación: ").bold = True
    p_m.add_run("31 de Julio de 2026\n")
    p_m.add_run("Auditor / Consultor: ").bold = True
    p_m.add_run("Antigravity (AI Coding & Web Consultant)")
    
    doc.add_paragraph() # Spacing

    # Helper for headings
    def add_custom_heading(text, level=1):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.bold = True
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = DARK_BLUE
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = GOLD
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
        return h

    # Section 1
    add_custom_heading("1. RESUMEN EJECUTIVO", level=1)
    p1 = doc.add_paragraph(
        "El sitio web danielsimons.xyz es un portal de marca personal y profesional alojado en la plataforma Blogger (Google). "
        "El sitio consolida una oferta diversificada de alto impacto profesional que abarca desde consultoría política y electoral "
        "(Campaña 2026 en Santa Cruz, Bolivia), desarrollo y evaluación de marca corporativa, propuestas inmobiliarias y urbanísticas, "
        "hasta herramientas educativas avanzadas como la 'Guía de Supervivencia a la Tesis' con un GPT Tutor integrado."
    )
    p1.paragraph_format.space_after = Pt(6)
    
    p2 = doc.add_paragraph(
        "El portal utiliza una plantilla personalizada basada en la estructura Soho de Blogger, con un tema de fondo oscuro y acentos dorados. "
        "Si bien cuenta con la estabilidad y seguridad que ofrece la infraestructura de Google, presenta amplias áreas de oportunidad en "
        "experiencia de usuario (UX), SEO técnico, optimización de conversión y automatización de contenidos."
    )
    p2.paragraph_format.space_after = Pt(10)

    # Section 2
    add_custom_heading("2. CRITERIOS DE EVALUACIÓN Y AUDITORÍA DETALLADA", level=1)
    
    add_custom_heading("A. Arquitectura Técnica y Plataforma", level=2)
    p_a = doc.add_paragraph()
    p_a.add_run("• CMS / Motor: ").bold = True
    p_a.add_run("Blogger (Google) con ID de Blog 433667097766389126.\n")
    p_a.add_run("• Servidor y CDN: ").bold = True
    p_a.add_run("Infraestructura global de Google con certificado SSL (HTTPS) activo y renovación automática.\n")
    p_a.add_run("• Sindicación de Contenidos: ").bold = True
    p_a.add_run("Feeds RSS 2.0 y Atom 1.0 habilitados en /feeds/posts/default.\n")
    p_a.add_run("• Diagnóstico Técnico: ").bold = True
    p_a.add_run("Excelente estabilidad y cero costo de servidor. Sin embargo, Blogger impone rigideces para aplicaciones dinámicas complejas.")
    
    add_custom_heading("B. SEO Técnico y Meta-información", level=2)
    p_b = doc.add_paragraph()
    p_b.add_run("• Indexabilidad y Canónicas: ").bold = True
    p_b.add_run("Canonical tag configurada hacia el dominio principal.\n")
    p_b.add_run("• Datos Estructurados: ").bold = True
    p_b.add_run("Implementación correcta de esquemas JSON-LD (BlogPosting, Person, Organization).\n")
    p_b.add_run("• Open Graph (Redes Sociales): ").bold = True
    p_b.add_run("Dispone de imágenes og:image en WebP optimizadas. ")
    p_b.add_run("Deficiencia crítica: ").bold = True
    p_b.add_run("La metaetiqueta og:description está actualmente vacía (content=''), afectando gravemente la vista previa al compartir en WhatsApp o LinkedIn.")

    add_custom_heading("C. Diseño Visual y Experiencia de Usuario (UI/UX)", level=2)
    p_c = doc.add_paragraph()
    p_c.add_run("• Paleta de Colores: ").bold = True
    p_c.add_run("Fondo negro (#000000), texto en blanco (#ffffff) y tonos dorados (#bca772). Ofrece elegancia pero puede generar contraste duro.\n")
    p_c.add_run("• Tipografía: ").bold = True
    p_c.add_run("Uso de EB Garamond (Serif) en el cuerpo y Montserrat/Lato en encabezados. Las fuentes Serif sobre fondo oscuro pueden causar fatiga visual en móviles.\n")
    p_c.add_run("• Navegación Móvil: ").bold = True
    p_c.add_run("Menú lateral desplegable y buscador integrados. Funcionamiento responsivo correcto.")

    add_custom_heading("D. Estrategia de Contenido y Conversión", level=2)
    p_d = doc.add_paragraph()
    p_d.add_run("• Diversidad Temática: ").bold = True
    p_d.add_run("Contenido de gran valor pero disperso entre política, branding, tesis y fotografía.\n")
    p_d.add_run("• Llamados a la Acción (CTA): ").bold = True
    p_d.add_run("Enlaces directos a WhatsApp y descargas de Brochure. Falta un formulario de captura de leads integrado directamente en la web.")

    # Section 3: FODA Table
    add_custom_heading("3. DIAGNÓSTICO FODA", level=1)
    
    table_foda = doc.add_table(rows=3, cols=2)
    table_foda.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    hdr0 = table_foda.cell(0, 0)
    hdr1 = table_foda.cell(0, 1)
    set_cell_background(hdr0, "1B365D")
    set_cell_background(hdr1, "1B365D")
    set_cell_margins(hdr0, 120, 120, 150, 150)
    set_cell_margins(hdr1, 120, 120, 150, 150)
    
    p = hdr0.paragraphs[0]
    r = p.add_run("FORTALEZAS (Internas)")
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    p = hdr1.paragraphs[0]
    r = p.add_run("OPORTUNIDADES (Externas)")
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # Row 1 (F & O)
    c_f = table_foda.cell(1, 0)
    c_o = table_foda.cell(1, 1)
    set_cell_background(c_f, "F9FAFC")
    set_cell_background(c_o, "F9FAFC")
    set_cell_margins(c_f, 100, 100, 120, 120)
    set_cell_margins(c_o, 100, 100, 120, 120)
    
    c_f.paragraphs[0].text = "• Servidor y SSL 100% gratuitos y estables (Google).\n• Innovación en propuestas (GPT Tutor, Data Electoral).\n• Marca personal sólida y estética elegante."
    c_o.paragraphs[0].text = "• Integrar un Asistente IA interactivo directamente en el sitio.\n• Crear Landing Pages dedicadas para conversión.\n• Automatizar artículos y newsletters mediante Blogger API."
    
    # Headers Row 2
    hdr2_0 = table_foda.cell(2, 0)
    hdr2_1 = table_foda.cell(2, 1)
    set_cell_background(hdr2_0, "555555")
    set_cell_background(hdr2_1, "555555")
    set_cell_margins(hdr2_0, 100, 100, 120, 120)
    set_cell_margins(hdr2_1, 100, 100, 120, 120)
    
    p = hdr2_0.paragraphs[0]
    r = p.add_run("DEBILIDADES (Internas)")
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    p = hdr2_1.paragraphs[0]
    r = p.add_run("AMENAZAS (Externas)")
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # Row 2 Content
    row3 = table_foda.add_row()
    c_d = row3.cells[0]
    c_a = row3.cells[1]
    set_cell_background(c_d, "FFFDF7")
    set_cell_background(c_a, "FFFDF7")
    set_cell_margins(c_d, 100, 100, 120, 120)
    set_cell_margins(c_a, 100, 100, 120, 120)
    
    c_d.paragraphs[0].text = "• Meta-descripciones Open Graph vacías.\n• Posible fatiga visual por tipografía Serif en fondo negro.\n• Ausencia de embudo de captación de clientes en la web."
    c_a.paragraphs[0].text = "• Pérdida de prospectos por depender únicamente de enlaces a WhatsApp.\n• Competidores con sitios dinámicos e interactivos de última generación."

    doc.add_paragraph() # Spacing

    # Section 4
    add_custom_heading("4. ¿CÓMO TE PUEDO AYUDAR?", level=1)
    
    p_h1 = doc.add_paragraph()
    p_h1.add_run("1. Rediseño y Modernización Frontend / UX:\n").bold = True
    p_h1.add_run("Optimización profunda del código XML de tu plantilla de Blogger o creación de una Landing Page de alta velocidad en React/Next.js/Vite integrada a tu dominio principal.\n")

    p_h2 = doc.add_paragraph()
    p_h2.add_run("2. Desarrollo de Asistentes e Integración de IA:\n").bold = True
    p_h2.add_run("Diseño de un Widget de Chatbot de IA flotante que atienda a los visitantes de danielsimons.xyz las 24 horas, respondiendo consultas sobre tus servicios de consultoría o guiando a los estudiantes con el GPT Tutor de tesis.\n")

    p_h3 = doc.add_paragraph()
    p_h3.add_run("3. SEO Técnico y Automatización de Metadatos:\n").bold = True
    p_h3.add_run("Corrección de metadatos vacíos, generación de fragmentos enriquecidos (Rich Snippets) y optimización de tarjetas para redes sociales.\n")

    p_h4 = doc.add_paragraph()
    p_h4.add_run("4. Dashboards e Informes Interactivos de Campaña / Transparencia:\n").bold = True
    p_h4.add_run("Desarrollo de gráficos e indicadores interactivos de datos electorales de Santa Cruz incrustables en tus publicaciones.\n")

    p_h5 = doc.add_paragraph()
    p_h5.add_run("5. Embudos de Conversión y Captura de Leads:\n").bold = True
    p_h5.add_run("Creación de formularios interactivos que envíen prospectos directamente a tu correo, CRM o WhatsApp de forma automatizada.")

    # Section 5
    add_custom_heading("5. ¿CÓMO ME PUEDO INTEGRAR A BLOGGER?", level=1)
    
    p_int = doc.add_paragraph(
        "Como asistente y desarrollador de IA, me puedo integrar con la plataforma Blogger de tu sitio mediante 4 modalidades clave:"
    )
    p_int.paragraph_format.space_after = Pt(6)

    tbl_int = doc.add_table(rows=5, cols=2)
    tbl_int.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    h0 = tbl_int.cell(0, 0)
    h1 = tbl_int.cell(0, 1)
    set_cell_background(h0, "1B365D")
    set_cell_background(h1, "1B365D")
    set_cell_margins(h0, 100, 100, 120, 120)
    set_cell_margins(h1, 100, 100, 120, 120)
    h0.paragraphs[0].add_run("MÉTODO DE INTEGRACIÓN").bold = True
    h0.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    h1.paragraphs[0].add_run("ALCANCE Y CASO DE USO").bold = True
    h1.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    methods = [
        ("1. Blogger REST API v3", "Conexión remota automatizada mediante scripts (Python/Node.js) para crear, programar y actualizar artículos, gestionar etiquetas y publicaciones sin entrar manualmente al panel de Blogger."),
        ("2. Widgets HTML/JS en Plantilla XML", "Inyección de scripts personalizados en el código fuente de Blogger. Permite incrustar Asistentes de IA flotantes, calculadoras, formularios de captura y elementos interactivos."),
        ("3. Flujo CI/CD Markdown -> Blogger", "Redacción de artículos en Markdown asistida por IA y conversión automática con formateo HTML limpio, imágenes optimizadas y metadatos SEO listos para ser publicados."),
        ("4. Sindicación vía RSS/Atom Feeds", "Lectura automatizada de tus entradas para alimentar boletines por correo electrónico, bots de Telegram/WhatsApp o resúmenes semanales de Inteligencia Artificial.")
    ]

    for idx, (m_title, m_desc) in enumerate(methods, start=1):
        cell_m = tbl_int.cell(idx, 0)
        cell_d = tbl_int.cell(idx, 1)
        bg = "F9FAFC" if idx % 2 != 0 else "FFFFFF"
        set_cell_background(cell_m, bg)
        set_cell_background(cell_d, bg)
        set_cell_margins(cell_m, 80, 80, 100, 100)
        set_cell_margins(cell_d, 80, 80, 100, 100)
        
        p = cell_m.paragraphs[0]
        p.add_run(m_title).bold = True
        cell_d.paragraphs[0].text = m_desc

    doc.add_paragraph() # Spacing

    # Section 6
    add_custom_heading("6. PLAN DE ACCIÓN RECOMENDADO", level=1)
    
    p_plan = doc.add_paragraph()
    p_plan.add_run("• Fase 1 (Inmediata - 24 a 48h): ").bold = True
    p_plan.add_run("Corregir metadatos SEO vacíos (og:description), mejorar el contraste tipográfico y optimizar llamados a la acción.\n")
    p_plan.add_run("• Fase 2 (Corto Plazo - 1 semana): ").bold = True
    p_plan.add_run("Inyectar el Widget de Asistente IA flotante en la plantilla XML de Blogger para atención automatizada a prospectos y tutoría de tesis.\n")
    p_plan.add_run("• Fase 3 (Mediano Plazo): ").bold = True
    p_plan.add_run("Implementar la API v3 de Blogger para automatización de contenidos y desarrollar dashboards interactivos para las propuestas electorales y comerciales.")

    # Save
    out_path = r"c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz\0_Informe_Auditoria_DanielSimons.docx"
    doc.save(out_path)
    print(f"Document saved to {out_path}")

if __name__ == "__main__":
    create_report()
