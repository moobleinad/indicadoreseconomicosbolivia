import os
import docx
from docx.shared import Pt, RGBColor, Inches

root_dir = r"c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz"
art_dir = os.path.join(root_dir, "02_ARTICULOS_Y_PUBLICACIONES")

md_path = os.path.join(art_dir, "02.03_Articulo_Serie_1_Resolucion_Ministerial_245.md")
docx_path = os.path.join(art_dir, "02.03_Articulo_Serie_1_Resolucion_Ministerial_245.docx")

content_md = """# RÉGIMEN CAMBIARIO: CÓMO FUNCIONA AHORA EL TIPO DE CAMBIO DEL DÓLAR EN BOLIVIA
### Serie Política Monetaria y Sociedad | Parte 1
### Por Daniel Simons

---

La **Resolución Ministerial N° 245** ha oficializado la transición hacia un **Régimen Cambiario Flexible** en Bolivia. El Estado ha dejado de fijar un precio congelado para el dólar, abriendo paso a un sistema donde la cotización se determina por la interacción diaria entre la oferta y la demanda de divisas en el sistema financiero.

Comprender cómo funciona este mecanismo y qué esperar a corto y mediano plazo es fundamental para resguardar el patrimonio personal y empresarial.

---

### 🏛️ 1. El mecanismo de oferta y demanda en el mercado monetario

El dinero funciona bajo la misma lógica que cualquier bien en el mercado:

- **Oferta de divisas:** Proviene de las exportaciones, el ingreso de remesas, los créditos internacionales y la venta de oro por parte del Banco Central de Bolivia (BCB).
- **Demanda de divisas:** Nace de los importadores de mercadería, repuestos y materia prima, así como de los ciudadanos que buscan proteger sus ahorros.

Bajo la R.M. 245, el BCB determina la cotización oficial registrando diariamente el punto de cruce entre esta oferta y demanda en el sistema bancario. Si la demanda de dólares supera a la oferta disponible, el valor de la divisa sube; si ingresan dólares al sistema, la cotización tiende a estabilizarse.

---

### 🔮 2. Proyección y presión inflacionaria: ¿Qué esperar para los próximos meses?

A corto plazo, la transición abrupta a un régimen flexible genera dos efectos inmediatos:

1. **Volatilidad continuada:** El tipo de cambio fluctuará mientras la economía busca su punto de equilibrio real.
2. **Presión inflacionaria (Inflación por costos):** El encarecimiento del dólar eleva el costo de los productos e insumos importados, trasladando esa presión de precios de forma directa a la canasta familiar y a la estructura operativa de los negocios.

---

### 🛡️ 3. Cómo resguardar tu dinero y el estímulo a la producción nacional

Frente a este escenario, existen tres estrategias fundamentales de resguardo y adaptación:

- **Resguardo en Dólares (Divisa Dura):** Pese a cualquier narrativa o discurso oficial, mantener reservas o activos indexados en moneda fuerte sigue siendo la vía principal de protección del capital frente a la devaluación del boliviano.
- **Refugio en Activos Físicos y Reales:** Convertir excedentes líquidos en activos tangibles —inventario no perecedero (arroz, alimentos, insumos básicos), inmuebles o terrenos— preserva el valor real del patrimonio frente al deterioro del poder adquisitivo.
- **Oportunidad para la Producción Nacional:** El encarecimiento de los productos importados genera un efecto de sustitución: los bienes producidos en Bolivia se vuelven más competitivos en precio frente a lo importado, lo que puede **estimular la industria y la producción local**, siempre que los productores adapten sus costos de insumos.

---

### 🔗 4. Conexión a la Parte 2: El impacto en el ciudadano

Este no es el mejor camino. Al igual que ocurrió con el ajuste repentino de los precios de los carburantes, este salto brusco coloca en aprietos directos al ciudadano común y le hace pagar a él los grandes ajustes acumulados de la economía.

¿Era la flotación libre inmediata la única opción, o un ajuste gradual (flotación sucia) hubiera protegido mejor al ciudadano y a la MYPE?

👉 **[IR A LA PARTE 2: El dilema de la flotación libre vs. el ajuste gradual en Bolivia]**

👉 **[EVALUAR ESTRUCTURA DE MI NEGOCIO CON DANIEL SIMONS](https://www.danielsimons.xyz/p/impulso-mype-360.html)**
"""

with open(md_path, "w", encoding="utf-8") as f:
    f.write(content_md)

# Save DOCX
doc = docx.Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p_title = doc.add_paragraph()
run_title = p_title.add_run("RÉGIMEN CAMBIARIO: CÓMO FUNCIONA AHORA EL TIPO DE CAMBIO DEL DÓLAR EN BOLIVIA")
run_title.font.name = "Montserrat"
run_title.font.size = Pt(16)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(188, 167, 114)

p_sub = doc.add_paragraph()
run_sub = p_sub.add_run("Serie Política Monetaria y Sociedad | Parte 1\nPor Daniel Simons")
run_sub.font.name = "Montserrat"
run_sub.font.size = Pt(11)
run_sub.font.bold = True
run_sub.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph("="*60)

for line in content_md.split("\n"):
    if line.startswith("### "):
        h = doc.add_heading(line.replace("### ", ""), level=3)
        for r in h.runs:
            r.font.name = "Montserrat"
            r.font.color.rgb = RGBColor(188, 167, 114)
    elif line.startswith("---") or line.startswith("# "):
        continue
    elif line.strip():
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Montserrat"
        r.font.size = Pt(10.5)

doc.save(docx_path)
print("SUCCESS: Final Series Article 1 created in MD and DOCX!")
