import os
import shutil
import docx
from docx.shared import Pt, RGBColor, Inches

root_dir = r"c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz"
master_md_path = os.path.join(root_dir, "00_DOCUMENTOS_ESTRATÉGICOS", "00.05_Prompt_Master_Diseno_y_Reglas_Web_DanielSimons.md")
master_docx_path = os.path.join(root_dir, "00_DOCUMENTOS_ESTRATÉGICOS", "00.05_Prompt_Master_Diseno_y_Reglas_Web_DanielSimons.docx")

# 1. UPDATE MASTER PROMPT WITH NO-EMOJI, NO-LINE AND DEDICATED ARTICLE FOLDER RULES
content_md = """# 📜 PROMPT MASTER & REGLAS INMUTABLES DE LA WEB
### Portal Oficial: [www.danielsimons.xyz](https://www.danielsimons.xyz/)

---

### 1. 🎯 POSICIONAMIENTO DE MARCA
- **Marca:** Daniel Simons — Estructurador de Ideas Complejas (IDEAS • ESTRUCTURA • IMPACTO).
- **Estilo Visual:** Ejecutivo, neutral, sobrio y elegante. Sin retórica política, sin gráficos infantiles y sin iconos estilo años 80.

---

### 2. 🎨 SISTEMA DE DISEÑO (REMBRANDT & CARAVAGGIO CHIAROSCURO)
- **Fondo Web:** Negro absoluto sólido (`#000000`) en toda la plataforma. Sin patrones ni texturas.
- **Paleta Oficial:**
  - **Fondo General y Secciones:** `#000000` (Negro Absoluto).
  - **Tarjetas:** `#0D0D0D` con marco fino dorado `rgba(188, 167, 114, 0.25)`.
  - **Detalles Dorados:** `#BCA772` (Dorado Cálido / Oro Envejecido).
  - **Texto Principal:** `#FFFFFF` y `#E0E0E0`.
- **Tipografía Móvil:** Títulos de tarjetas en celulares fijados en **`15px !important` (Negrita)**.

---

### 3. 📐 ESTRUCTURA DE LA PORTADA (6 SECCIONES EN CARRUSEL HORIZONTAL)
La portada se compone de filas con deslizamiento horizontal en carrusel (flechas doradas `❮` `❯`):

1. **ÚLTIMAS ENTRADAS DEL BLOG:** Carrusel de artículos con botón *"VER TODAS ➔"*.
2. **SERVICIOS DE ESTRUCTURACIÓN:** *Destilado de Ideas*, *FORJA*, *Marketing 360°*, *Impulso MYPE*.
3. **PROYECTOS PROPIOS & INVESTIGACIÓN:** *Modelo MFEIR*, *El Juego del Emprendedor*.
4. **LIBROS:** Sección dedicada a publicaciones y obras de Daniel Simons.
5. **QUIÉN SOY (DANIEL SIMONS):** Tarjeta ejecutiva de presentación.
6. **TRABAJOS Y PROPUESTAS:** *Propuesta Urbanizaciones*, *Transparencia Inteligente*, *Guía Sobreviviendo a la Tesis*, *Evaluar Desarrollo de Marca*.

---

### 4. 🧠 REGLAS DE REDACCIÓN, ESTÍMULOS SUPERNORMALES Y CERO EMOJIS / LÍNEAS
- **PROHIBICIÓN ABSOLUTA DE ICONOS, EMOJIS Y LÍNEAS SEPARADORAS (`---`):**
  * NUNCA JAMÁS colocar emojis ni iconos en títulos, subtítulos o cuerpo de texto.
  * NUNCA JAMÁS colocar rayas separadoras de texto (`---` / `<hr/>`) dentro del cuerpo de los artículos.
- **ENFOQUE GERENCIAL Y AHORRO DE TIEMPO:** Redacción ultra-práctica, directa al grano, utilitaria, concisa y quirúrgicamente ejecutiva.
- **ESTÍMULOS SUPERNORMALES (NEURO-COGNICIÓN):** Títulos quirúrgicos, máxima sustancia técnica por línea y ganchos de conversión inmediata para maximizar la posibilidad de que el lector contacte por WhatsApp o pase al siguiente artículo/servicio.

---

### 5. ✍️ METODOLOGÍA OFICIAL DE PROPUESTA, SELECCIÓN, REDACCIÓN Y APROBACIÓN DE ARTÍCULOS
1. **Propuesta de Tendencias:** El asistente presenta los temas de mayor tendencia coyuntural.
2. **Selección del Tema por Daniel:** Daniel elige cuál de las opciones desarrollaremos.
3. **Presentación de Borrador Directo e Imagen:** El asistente redacta el artículo bajo normas de **Neuro-Redacción Utilitaria** (sin emojis ni líneas separadoras) y entrega la imagen `.webp` oficial.
4. **Revisión, Ajustes y Aprobación de Daniel:** Daniel revisa el texto y la foto, solicita correcciones y otorga su aprobación explícita.
5. **Publicación en Blogger:** Daniel sube el artículo aprobado e inserta la imagen nativamente desde el editor de Blogger.
6. **Vincular en Portada:** Daniel comparte el enlace publicado y el asistente vincula la URL oficial de la imagen en la miniatura cuadrada (1:1) de la portada.

---

### 6. 📁 ESTRUCTURA DE ARTÍCULOS EN CARPETAS DEDICADAS Y SISTEMA NUMÉRICO
- **CARPETAS INDIVIDUALES POR ARTÍCULO:** Cada artículo se organizará dentro de su propia carpeta dedicada en `02_ARTICULOS_Y_PUBLICACIONES/` conteniendo exclusivamente sus 3 archivos:
  - Archivo `.md` (Limpio sin emojis ni líneas `---`)
  - Archivo `.docx` (Microsoft Word)
  - Imagen `.webp` oficial optimizada
- **Jerarquía Numérica:** Carpetas y archivos organizados con prefijo numérico secuencial (`00_...`, `01_...`, `02_...`).

---

### 7. ⚡ LÍMITES TÉCNICOS Y PESO EN BLOGGER
- **Peso Máximo del XML:** El archivo XML del tema NUNCA DEBE SUPERAR LOS 200 KB (Rango ideal: 160 KB – 190 KB).
- **CLÁUSULA DE CONTROL ABSOLUTO:** Sólo modificarás LO QUE EXPRESAMENTE SE TE INSTRUYA; TODO LO DEMÁS QUEDARÁ EXACTAMENTE IGUAL.
"""

with open(master_md_path, "w", encoding="utf-8") as f:
    f.write(content_md)

# Save Master DOCX
doc = docx.Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p_title = doc.add_paragraph()
run_title = p_title.add_run("PROMPT MASTER & REGLAS INMUTABLES DE LA WEB")
run_title.font.name = "Montserrat"
run_title.font.size = Pt(20)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(188, 167, 114)

p_sub = doc.add_paragraph()
run_sub = p_sub.add_run("Portal Oficial: www.danielsimons.xyz | Daniel Simons")
run_sub.font.name = "Montserrat"
run_sub.font.size = Pt(11)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph("="*60)

def add_sec(title, text_lines):
    h = doc.add_heading(title, level=2)
    for r in h.runs:
        r.font.name = "Montserrat"
        r.font.color.rgb = RGBColor(188, 167, 114)
    for line in text_lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Montserrat"
        r.font.size = Pt(10.5)

add_sec("1. POSICIONAMIENTO DE MARCA", [
    "• Marca: Daniel Simons — Estructurador de Ideas Complejas (IDEAS • ESTRUCTURA • IMPACTO).",
    "• Estilo Visual: Ejecutivo, neutral, sobrio y elegante. Sin retórica política, sin gráficos infantiles y sin iconos estilo años 80."
])

add_sec("2. SISTEMA DE DISEÑO (REMBRANDT & CARAVAGGIO CHIAROSCURO)", [
    "• Fondo Web: Negro absoluto sólido (#000000) en toda la plataforma. Sin patrones ni texturas.",
    "• Paleta Oficial: Fondo General (#000000), Tarjetas (#0D0D0D) con marco dorado (rgba(188, 167, 114, 0.25)), Detalles Dorados (#BCA772), Texto (#FFFFFF / #E0E0E0).",
    "• Tipografía Móvil: Títulos de tarjetas en celulares fijados en 15px !important (Negrita)."
])

add_sec("3. ESTRUCTURA DE LA PORTADA (CARRUSEL HORIZONTAL)", [
    "1. ÚLTIMAS ENTRADAS DEL BLOG: Carrusel de artículos con botón 'VER TODAS ➔'.",
    "2. SERVICIOS DE ESTRUCTURACIÓN: Destilado de Ideas, FORJA, Marketing 360°, Impulso MYPE.",
    "3. PROYECTOS PROPIOS & INVESTIGACIÓN: Modelo MFEIR, El Juego del Emprendedor.",
    "4. LIBROS: Sección oficial de libros y publicaciones.",
    "5. QUIÉN SOY (DANIEL SIMONS): Tarjeta ejecutiva de presentación.",
    "6. TRABAJOS Y PROPUESTAS: Urbanizaciones, Transparencia Inteligente, Guía Tesis, Evaluar Marca."
])

add_sec("4. REGLAS DE REDACCIÓN, ESTÍMULOS SUPERNORMALES Y CERO EMOJIS / LÍNEAS", [
    "• PROHIBICIÓN ABSOLUTA DE ICONOS, EMOJIS Y LÍNEAS SEPARADORAS (---): NUNCA JAMÁS colocar emojis ni iconos en títulos o cuerpos de texto. NUNCA JAMÁS colocar rayas separadoras de texto en artículos.",
    "• ENFOQUE GERENCIAL Y AHORRO DE TIEMPO: Redacción ultra-práctica, directa al grano, utilitaria y quirúrgicamente ejecutiva.",
    "• ESTÍMULOS SUPERNORMALES (NEURO-COGNICIÓN): Títulos quirúrgicos, máxima sustancia técnica por línea y ganchos de conversión directa hacia WhatsApp, servicios o el siguiente artículo."
])

add_sec("5. METODOLOGÍA OFICIAL DE PROPUESTA, SELECCIÓN, REDACCIÓN Y APROBACIÓN DE ARTÍCULOS", [
    "1. Propuesta de Tendencias: Presentación de temas coyunturales.",
    "2. Selección del Tema por Daniel: Elección del tema a desarrollar.",
    "3. Presentación de Borrador Directo e Imagen: Redacción bajo normas de Neuro-Redacción Utilitaria (sin emojis ni líneas separadoras) e imagen .webp.",
    "4. Revisión y Aprobación de Daniel: Revisión, ajustes y aprobación explícita.",
    "5. Publicación en Blogger: Subida manual con imagen nativa.",
    "6. Vincular en Portada: Vinculación de la URL oficial en la miniatura cuadrada 1:1."
])

add_sec("6. ESTRUCTURA DE ARTÍCULOS EN CARPETAS DEDICADAS Y SISTEMA NUMÉRICO", [
    "• CARPETAS INDIVIDUALES POR ARTÍCULO: Cada artículo se organizará dentro de su propia carpeta dedicada en 02_ARTICULOS_Y_PUBLICACIONES/ conteniendo exclusivamente sus 3 archivos: archivo .md (sin emojis/líneas), archivo .docx y su imagen .webp oficial.",
    "• Jerarquía Numérica: Carpetas y archivos organizados con prefijo numérico secuencial (00_..., 01_...)."
])

add_sec("7. LÍMITES TÉCNICOS Y PESO EN BLOGGER", [
    "• Peso Máximo del XML: NUNCA DEBE SUPERAR LOS 200 KB (Rango ideal: 160 KB – 190 KB).",
    "• CLÁUSULA DE CONTROL ABSOLUTO: Sólo modificarás LO QUE EXPRESAMENTE SE TE INSTRUYA; TODO LO DEMÁS QUEDARÁ EXACTAMENTE IGUAL."
])

doc.save(master_docx_path)
print("SUCCESS: Updated Master MD and DOCX with No-Emoji/No-Line & Dedicated Article Folders Rule!")

# 2. RE-STRUCTURE ARTICLE DIRECTORIES IN 02_ARTICULOS_Y_PUBLICACIONES/
art_base_dir = os.path.join(root_dir, "02_ARTICULOS_Y_PUBLICACIONES")

# Article 1 BTH
dir_art1 = os.path.join(art_base_dir, "02.01_Articulo_1_El_Drama_del_BTH_DanielSimons")
os.makedirs(dir_art1, exist_ok=True)
for item in ["02.01_Articulo_1_El_Drama_del_BTH_DanielSimons.md", "02.01_Articulo_1_El_Drama_del_BTH_DanielSimons.docx"]:
    p = os.path.join(art_base_dir, item)
    if os.path.exists(p):
        shutil.move(p, os.path.join(dir_art1, item))
img_art1 = os.path.join(root_dir, "01_LINEA_GRAFICA_Y_ASSETS", "01.02_foto_articulo1_bth_oficial.webp")
if os.path.exists(img_art1):
    shutil.copy(img_art1, os.path.join(dir_art1, "02.01_foto_articulo1_bth_oficial.webp"))

# Article 2 TikTok/MYPE
dir_art2 = os.path.join(art_base_dir, "02.02_Articulo_2_Por_Que_TikTok_No_Salvara_Tu_MYPE")
os.makedirs(dir_art2, exist_ok=True)
for item in ["02.02_Articulo_2_Por_Que_TikTok_No_Salvara_Tu_MYPE.md", "02.02_Articulo_2_Por_Que_TikTok_No_Salvara_Tu_MYPE.docx"]:
    p = os.path.join(art_base_dir, item)
    if os.path.exists(p):
        shutil.move(p, os.path.join(dir_art2, item))
img_art2 = os.path.join(root_dir, "01_LINEA_GRAFICA_Y_ASSETS", "01.06_foto_articulo2_dolar_mype_oficial.webp")
if os.path.exists(img_art2):
    shutil.copy(img_art2, os.path.join(dir_art2, "02.02_foto_articulo2_dolar_mype_oficial.webp"))

# Article Series 1 RM 245
dir_art_rm245 = os.path.join(art_base_dir, "02.03_Articulo_Serie_1_Resolucion_Ministerial_245")
os.makedirs(dir_art_rm245, exist_ok=True)

# Generate CLEAN MD & DOCX for Series Article 1 (ZERO EMOJIS, ZERO '---' LINES!)
clean_md_content = """# RÉGIMEN CAMBIARIO: CÓMO FUNCIONA AHORA EL TIPO DE CAMBIO DEL DÓLAR EN BOLIVIA
### Serie Política Monetaria y Sociedad | Parte 1
### Por Daniel Simons

La Resolución Ministerial N° 245 ha oficializado la transición hacia un Régimen Cambiario Flexible en Bolivia. El Estado ha dejado de fijar un precio congelado para el dólar, abriendo paso a un sistema donde la cotización se determina por la interacción diaria entre la oferta y la demanda de divisas en el sistema financiero.

Comprender cómo funciona este mecanismo y qué esperar a corto y mediano plazo es fundamental para resguardar el patrimonio personal y empresarial.

1. El mecanismo de oferta y demanda en el mercado monetario

El dinero funciona bajo la misma lógica que cualquier bien en el mercado:

Oferta de divisas: Proviene de las exportaciones, el ingreso de remesas, los créditos internacionales y la venta de oro por parte del Banco Central de Bolivia (BCB).
Demanda de divisas: Nace de los importadores de mercadería, repuestos y materia prima, así como de los ciudadanos que buscan proteger sus ahorros.

Bajo la R.M. 245, el BCB determina la cotización oficial registrando diariamente el punto de cruce entre esta oferta y demanda en el sistema bancario. Si la demanda de dólares supera a la oferta disponible, el valor de la divisa sube; si ingresan dólares al sistema, la cotización tiende a estabilizarse.

2. Proyección y presión inflacionaria: ¿Qué esperar para los próximos meses?

A corto plazo, la transición abrupta a un régimen flexible genera dos efectos inmediatos:

1. Volatilidad continuada: El tipo de cambio fluctuará mientras la economía busca su punto de equilibrio real.
2. Presión inflacionaria (Inflación por costos): El encarecimiento del dólar eleva el costo de los productos e insumos importados, trasladando esa presión de precios de forma directa a la canasta familiar y a la estructura operativa de los negocios.

3. Cómo resguardar tu dinero y el estímulo a la producción nacional

Frente a este escenario, existen tres estrategias fundamentales de resguardo y adaptación:

Resguardo en Dólares (Divisa Dura): Pese a cualquier narrativa o discurso oficial, mantener reservas o activos indexados en moneda fuerte sigue siendo la vía principal de protección del capital frente a la devaluación del boliviano.
Refugio en Activos Físicos y Reales: Convertir excedentes líquidos en activos tangibles —inventario no perecedero (arroz, alimentos, insumos básicos), inmuebles o terrenos— preserva el valor real del patrimonio frente al deterioro del poder adquisitivo.
Oportunidad para la Producción Nacional: El encarecimiento de los productos importados genera un efecto de sustitución: los bienes producidos en Bolivia se vuelven más competitivos en precio frente a lo importado, lo que puede estimular la industria y la producción local, siempre que los productores adapten sus costos de insumos.

4. Conexión a la Parte 2: El impacto en el ciudadano

Este no es el mejor camino. Al igual que ocurrió con el ajuste repentino de los precios de los carburantes, este salto brusco coloca en aprietos directos al ciudadano común y le hace pagar a él los grandes ajustes acumulados de la economía.

¿Era la flotación libre inmediata la única opción, o un ajuste gradual (flotación sucia) hubiera protegido mejor al ciudadano y a la MYPE?

[IR A LA PARTE 2: El dilema de la flotación libre vs. el ajuste gradual en Bolivia]

[EVALUAR ESTRUCTURA DE MI NEGOCIO CON DANIEL SIMONS](https://www.danielsimons.xyz/p/impulso-mype-360.html)
"""

clean_md_path = os.path.join(dir_art_rm245, "02.03_Articulo_Serie_1_Resolucion_Ministerial_245.md")
clean_docx_path = os.path.join(dir_art_rm245, "02.03_Articulo_Serie_1_Resolucion_Ministerial_245.docx")

with open(clean_md_path, "w", encoding="utf-8") as f:
    f.write(clean_md_content)

doc_c = docx.Document()
for section in doc_c.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p_title = doc_c.add_paragraph()
run_title = p_title.add_run("RÉGIMEN CAMBIARIO: CÓMO FUNCIONA AHORA EL TIPO DE CAMBIO DEL DÓLAR EN BOLIVIA")
run_title.font.name = "Montserrat"
run_title.font.size = Pt(16)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(188, 167, 114)

p_sub = doc_c.add_paragraph()
run_sub = p_sub.add_run("Serie Política Monetaria y Sociedad | Parte 1\nPor Daniel Simons")
run_sub.font.name = "Montserrat"
run_sub.font.size = Pt(11)
run_sub.font.bold = True
run_sub.font.color.rgb = RGBColor(128, 128, 128)

for line in clean_md_content.split("\n"):
    if line.startswith("# "):
        continue
    elif line.strip():
        p = doc_c.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Montserrat"
        r.font.size = Pt(10.5)

doc_c.save(clean_docx_path)

img_rm245 = os.path.join(root_dir, "01_LINEA_GRAFICA_Y_ASSETS", "01.07_foto_articulo1_rm245_oficial.webp")
if os.path.exists(img_rm245):
    shutil.copy(img_rm245, os.path.join(dir_art_rm245, "02.03_foto_articulo1_rm245_oficial.webp"))

print("SUCCESS: Cleaned Article 1 Series and organized all dedicated article folders!")
