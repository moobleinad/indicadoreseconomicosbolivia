import os
import docx

md_art1_path = r'c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz\0_Articulo_1_El_Drama_del_BTH_DanielSimons.md'
docx_art1_path = r'c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz\0_Articulo_1_El_Drama_del_BTH_DanielSimons.docx'

art1_clean = """# De la idea escolar al proyecto ordenado: Cómo estructurar emprendimientos juveniles con claridad e impacto real

**Por Daniel Simons**  
*Estructurador de Ideas Complejas | Autor de «El Juego del Emprendedor»*

---

### [ATENCIÓN] El verdadero reto de los proyectos de emprendimiento escolar

Cada año, miles de estudiantes de secundaria y jóvenes emprendedores se enfrentan al desafío de presentar un proyecto productivo o iniciativa de negocio. Tanto para los alumnos como para sus familias y docentes, la primera gran barrera no es la falta de creatividad o entusiasmo, sino una pregunta fundamental: **¿por dónde se empieza y cómo se pasa de una idea en la cabeza a una propuesta clara y ordenada?**

En la mayoría de las ferias y materias prácticas, es común ver un esfuerzo enorme enfocado únicamente en la presentación visual del producto. Sin embargo, cuando se le pregunta al estudiante sobre la utilidad real de su idea, sus costos o su público objetivo, surge la confusión. 

---

### [INTERÉS] El valor del orden: Transformar el desorden en estructura lógica

Emprender en la etapa escolar o universitaria no requiere presentar de inmediato una empresa gigante ni buscar financiamiento bancario a gran escala. Ser realistas y honestos es el primer paso: **un proyecto juvenil no necesita ser perfecto ni financiable de inmediato para ser valioso**.

Lo que verdaderamente transforma la experiencia de un estudiante es adquirir **orden, criterio técnico y estructura**. 

Pasar del caos a la claridad implica responder cuatro preguntas clave de forma sencilla:
1. **El Problema Real:** ¿Qué necesidad o deseo concreto busca resolver este proyecto?
2. **El Público Objetivo:** ¿A quién le interesa realmente este producto o servicio y por qué?
3. **Los Costos Reales:** ¿Cuánto cuesta producirlo y a qué precio se debe ofrecer sin trabajar a pérdida?
4. **La Propuesta Clara:** ¿Cómo se presenta la idea de forma lógica, comprensible y honesta?

---

### [DESEO] Una expectativa honesta: Claridad antes que ilusión

El objetivo de trabajar con métodos estructurados de emprendimiento no es prometer resultados irreales o bancables de la noche a la mañana. 

La verdadera ganancia para el joven, el padre de familia y el docente radica en dar **un paso sólido hacia adelante**:
* **Del desorden a la estructura:** El estudiante deja de improvisar y aprende a pensar con lógica de negocio.
* **De la duda a la seguridad:** Comprende los números y los fundamentos de su trabajo, lo que le permite defender su idea con confianza ante cualquier evaluación.
* **De la teoría a la práctica útil:** Desarrolla un producto mínimo ordenado y entendible que sienta las bases para su futuro profesional.

Aunque el proyecto aún esté lejos de un financiamiento comercial, haber logrado orden, claridad y criterio es el mayor activo que un estudiante puede llevarse para la vida real.

---

### [ACCIÓN] Construyendo claridad paso a paso

Si eres docente, director o padre de familia y deseas que tus jóvenes aprendan a transformar sus ideas en proyectos ordenados y con sentido práctico:

* 📘 **«El Juego del Emprendedor»:** La guía metodológica simplificada diseñada para acompañar a estudiantes y jóvenes en el diseño paso a paso de sus proyectos.
* 🎓 **Talleres & Estructuración Guiada:** Capacitaciones prácticas para colegios, institutos y universidades.

📲 **Contacto directo con Daniel Simons:** [WhatsApp de Consultoría y Estructuración](https://wa.me/59170000000)  
🌐 **Portal Oficial:** [www.danielsimons.xyz](https://www.danielsimons.xyz)
"""

with open(md_art1_path, 'w', encoding='utf-8') as f:
    f.write(art1_clean)

print("SUCCESS: 0_Articulo_1_El_Drama_del_BTH_DanielSimons.md updated successfully!")

try:
    doc = docx.Document()
    doc.add_heading("De la idea escolar al proyecto ordenado", 0)
    doc.add_paragraph("Por Daniel Simons - Estructurador de Ideas Complejas")
    doc.add_paragraph("-" * 50)

    for line in art1_clean.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.strip():
            doc.add_paragraph(line.strip())

    doc.save(docx_art1_path)
    print("SUCCESS: 0_Articulo_1_El_Drama_del_BTH_DanielSimons.docx updated successfully!")
except Exception as e:
    print("WARN: Could not write docx file (file may be open in Word):", e)
