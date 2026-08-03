import os
import docx
from docx.shared import Pt, RGBColor, Inches

root_dir = r"c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz"
master_md_path = os.path.join(root_dir, "00_DOCUMENTOS_ESTRATÉGICOS", "00.05_Prompt_Master_Diseno_y_Reglas_Web_DanielSimons.md")
master_docx_path = os.path.join(root_dir, "00_DOCUMENTOS_ESTRATÉGICOS", "00.05_Prompt_Master_Diseno_y_Reglas_Web_DanielSimons.docx")

content_md = """# 📜 PROMPT MASTER & REGLAS INMUTABLES DE LA WEB
### Portal Oficial: [www.danielsimons.xyz](https://www.danielsimons.xyz/)

---

### 1. 🎯 POSICIONAMIENTO DE MARCA
- **Marca:** Daniel Simons — Estructurador de Ideas Complejas (IDEAS • ESTRUCTURA • IMPACTO).
- **Estilo Visual:** Ejecutivo, neutral, sobrio y elegante. Sin retórica política, sin gráficos infantiles y sin iconos estilo años 80.

---

### 2. 🎨 SISTEMA DE DISEÑO (REMBRANDT & CARAVAGGIO CHIAROSCURO)
- **Fondo Web:** Negro absoluto sólido (`#000000`) en toda la plataforma. Sin patrones ni texturas.
- **Paleta Oficial:**
  - **Fondo General y Secciones:** `#000000` (Negro Absoluto).
  - **Tarjetas:** `#0D0D0D` con marco fino dorado `rgba(188, 167, 114, 0.25)`.
  - **Detalles Dorados:** `#BCA772` (Dorado Cálido / Oro Envejecido).
  - **Texto Principal:** `#FFFFFF` y `#E0E0E0`.
- **Tipografía Móvil:** Títulos de tarjetas en celulares fijados en **`15px !important` (Negrita)**.

---

### 3. 📐 ESTRUCTURA DE LA PORTADA (6 SECCIONES EN CARRUSEL HORIZONTAL)
La portada se compone de filas con deslizamiento horizontal en carrusel (flechas doradas `❮` `❯`):

1. **ÚLTIMAS ENTRADAS DEL BLOG:** Carrusel de artículos con botón *"VER TODAS ➔"*.
2. **SERVICIOS DE ESTRUCTURACIÓN:** *Destilado de Ideas*, *FORJA*, *Marketing 360°*, *Impulso MYPE*.
3. **PROYECTOS PROPIOS & INVESTIGACIÓN:** *Modelo MFEIR*, *El Juego del Emprendedor*.
4. **LIBROS:** Sección dedicada a publicaciones y obras de Daniel Simons.
5. **QUIÉN SOY (DANIEL SIMONS):** Tarjeta ejecutiva de presentación.
6. **TRABAJOS Y PROPUESTAS:** *Propuesta Urbanizaciones*, *Transparencia Inteligente*, *Guía Sobreviviendo a la Tesis*, *Evaluar Desarrollo de Marca*.

---

### 4. 📷 REGLAS DE IMÁGENES Y MÉTODO DE VINCULACIÓN CDN
- **Formato:** Miniaturas **100% cuadradas (1:1)** que ocupan toda la cabecera de la tarjeta (`object-fit: cover`).
- **MÉTODO ESTÁNDAR INMUTABLE PARA PÁGINAS:**
  Para cualquier cambio o adición a las páginas de *Servicios*, *Proyectos Propios*, *Trabajos y Propuestas* o cualquier página que se muestre en la portada principal, **SE UTILIZARÁ EXCLUSIVAMENTE EL MÉTODO DE VINCULACIÓN DIRECTA DESDE EL CDN DE BLOGGER (`blogger.googleusercontent.com`)**.

---

### 5. ✍️ METODOLOGÍA OFICIAL DE PROPUESTA, SELECCIÓN, REDACCIÓN Y APROBACIÓN DE ARTÍCULOS
La metodología estricta a seguir para cada nuevo artículo será:

1. **Propuesta de Tendencias:** El asistente presenta los temas de mayor tendencia coyuntural.
2. **Selección del Tema por Daniel:** Daniel elige cuál de las opciones desarrollaremos (ej. Opción 2: Dólar y MYPEs).
3. **Presentación de Borrador e Imagen:** El asistente redacta el artículo completo bajo la **estructura AIDA** y genera la imagen oficial optimizada en formato `.webp`, entregándolos para revisión.
4. **Revisión, Ajustes y Aprobación de Daniel:** Daniel revisa el texto y la foto, solicita correcciones si existen, y otorga su aprobación explícita.
5. **Publicación en Blogger:** Daniel sube el artículo aprobado e inserta la imagen nativamente desde el editor de Blogger.
6. **Vincular en Portada:** Daniel comparte el enlace publicado y el asistente vincula la URL oficial de la imagen en la miniatura cuadrada (1:1) del carrusel de la portada.

---

### 6. 📁 SISTEMA NUMÉRICO Y NORMAS DE ORGANIZACIÓN DE ARCHIVOS
- **Jerarquía Numérica:** Carpetas y archivos organizados con prefijo numérico secuencial (`00_...`, `01_...`, `02_...`).
- **Subdivisión Decimal:** Documentos internos numerados decimalmente (`00.01_...`, `00.02_...`).

---

### 7. ⚡ LÍMITES TÉCNICOS Y PESO EN BLOGGER
- **Peso Máximo del XML:** El archivo XML del tema NUNCA DEBE SUPERAR LOS 200 KB (Rango ideal: 160 KB – 190 KB).
- **CLÁUSULA DE CONTROL ABSOLUTO:** Sólo modificarás LO QUE EXPRESAMENTE SE TE INSTRUYA; TODO LO DEMÁS QUEDARÁ EXACTAMENTE IGUAL.
"""

with open(master_md_path, "w", encoding="utf-8") as f:
    f.write(content_md)

# Save Master DOCX
doc = docx.Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p_title = doc.add_paragraph()
run_title = p_title.add_run("PROMPT MASTER & REGLAS INMUTABLES DE LA WEB")
run_title.font.name = "Montserrat"
run_title.font.size = Pt(20)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(188, 167, 114)

p_sub = doc.add_paragraph()
run_sub = p_sub.add_run("Portal Oficial: www.danielsimons.xyz | Daniel Simons")
run_sub.font.name = "Montserrat"
run_sub.font.size = Pt(11)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph("="*60)

def add_sec(title, text_lines):
    h = doc.add_heading(title, level=2)
    for r in h.runs:
        r.font.name = "Montserrat"
        r.font.color.rgb = RGBColor(188, 167, 114)
    for line in text_lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Montserrat"
        r.font.size = Pt(10.5)

add_sec("1. POSICIONAMIENTO DE MARCA", [
    "• Marca: Daniel Simons — Estructurador de Ideas Complejas (IDEAS • ESTRUCTURA • IMPACTO).",
    "• Estilo Visual: Ejecutivo, neutral, sobrio y elegante. Sin retórica política, sin gráficos infantiles y sin iconos estilo años 80."
])

add_sec("2. SISTEMA DE DISEÑO (REMBRANDT & CARAVAGGIO CHIAROSCURO)", [
    "• Fondo Web: Negro absoluto sólido (#000000) en toda la plataforma. Sin patrones ni texturas.",
    "• Paleta Oficial: Fondo General (#000000), Tarjetas (#0D0D0D) con marco dorado (rgba(188, 167, 114, 0.25)), Detalles Dorados (#BCA772), Texto (#FFFFFF / #E0E0E0).",
    "• Tipografía Móvil: Títulos de tarjetas en celulares fijados en 15px !important (Negrita)."
])

add_sec("3. ESTRUCTURA DE LA PORTADA (CARRUSEL HORIZONTAL)", [
    "1. ÚLTIMAS ENTRADAS DEL BLOG: Carrusel de artículos con botón 'VER TODAS ➔'.",
    "2. SERVICIOS DE ESTRUCTURACIÓN: Destilado de Ideas, FORJA, Marketing 360°, Impulso MYPE.",
    "3. PROYECTOS PROPIOS & INVESTIGACIÓN: Modelo MFEIR, El Juego del Emprendedor.",
    "4. LIBROS: Sección oficial de libros y publicaciones.",
    "5. QUIÉN SOY (DANIEL SIMONS): Tarjeta ejecutiva de presentación.",
    "6. TRABAJOS Y PROPUESTAS: Urbanizaciones, Transparencia Inteligente, Guía Tesis, Evaluar Marca."
])

add_sec("4. REGLAS DE IMÁGENES Y MÉTODO DE VINCULACIÓN CDN", [
    "• Formato: Miniaturas 100% cuadradas (1:1) llenando la tarjeta (object-fit: cover).",
    "• MÉTODO ESTÁNDAR INMUTABLE: Para cualquier modificación o adición en las páginas de Servicios, Proyectos o Trabajos que vayan en la portada principal, SE UTILIZARÁ EXCLUSIVAMENTE EL MÉTODO DE VINCULACIÓN DIRECTA DESDE EL CDN DE BLOGGER (blogger.googleusercontent.com)."
])

add_sec("5. METODOLOGÍA OFICIAL DE PROPUESTA, SELECCIÓN, REDACCIÓN Y APROBACIÓN DE ARTÍCULOS", [
    "1. Propuesta de Tendencias: El asistente presenta los temas de mayor tendencia coyuntural.",
    "2. Selección del Tema por Daniel: Daniel elige cuál de las opciones desarrollaremos (ej. Opción 2: Dólar y MYPEs).",
    "3. Presentación de Borrador e Imagen: El asistente redacta el artículo completo bajo la estructura AIDA y genera la imagen oficial optimizada en formato .webp, entregándolos para revisión.",
    "4. Revisión, Ajustes y Aprobación de Daniel: Daniel revisa el texto y la foto, solicita correcciones si existen, y otorga su aprobación explícita.",
    "5. Publicación en Blogger: Daniel sube el artículo aprobado e inserta la imagen nativamente desde el editor de Blogger.",
    "6. Vincular en Portada: Daniel comparte el enlace publicado y el asistente vincula la URL oficial de la imagen en la miniatura cuadrada (1:1) del carrusel de la portada."
])

add_sec("6. SISTEMA NUMÉRICO Y NORMAS DE ORGANIZACIÓN DE ARCHIVOS", [
    "• Jerarquía Numérica: Carpetas y archivos organizados con prefijo numérico secuencial (00_..., 01_..., 02_...).",
    "• Subdivisión Decimal: Documentos internos numerados decimalmente (00.01_..., 00.02_...)."
])

add_sec("7. LÍMITES TÉCNICOS Y PESO EN BLOGGER", [
    "• Peso Máximo del XML: NUNCA DEBE SUPERAR LOS 200 KB (Rango ideal: 160 KB – 190 KB).",
    "• CLÁUSULA DE CONTROL ABSOLUTO: Sólo modificarás LO QUE EXPRESAMENTE SE TE INSTRUYA; TODO LO DEMÁS QUEDARÁ EXACTAMENTE IGUAL."
])

doc.save(master_docx_path)
print("SUCCESS: Updated Master MD and DOCX with Article Selection & Approval Workflow Rule!")
