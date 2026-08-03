import os
import docx
from docx.shared import Pt, RGBColor, Inches
from PIL import Image

root_dir = r"c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz"
md_path = os.path.join(root_dir, "0_Prompt_Master_Diseno_y_Reglas_Web_DanielSimons.md")
docx_path = os.path.join(root_dir, "0_Prompt_Master_Diseno_y_Reglas_Web_DanielSimons.docx")

# STEP 1: CONVERT ARTICLE 1 IMAGE TO OPTIMIZED WEBP
source_img_path = r"C:\Users\Usuario\.gemini\antigravity\brain\953edcf7-5053-4ba6-94e5-82edfed7e1d2\foto_articulo1_bth_v4_1785592381896.jpg"
out_webp_path = os.path.join(root_dir, "foto_articulo1_bth_oficial.webp")

im = Image.open(source_img_path).convert("RGB")
im.save(out_webp_path, "WEBP", quality=80)
print(f"Generated WebP: {out_webp_path} (Size: {os.path.getsize(out_webp_path)/1024:.2f} KB)")

# STEP 2: UPDATE MASTER PROMPT MD & DOCX WITH THE WORKFLOW INSTRUCTION
content_md = """# 📜 PROMPT MASTER & REGLAS INMUTABLES DE LA WEB
### Portal Oficial: [www.danielsimons.xyz](https://www.danielsimons.xyz/)

---

### 1. 🎯 POSICIONAMIENTO DE MARCA
- **Marca:** Daniel Simons — Estructurador de Ideas Complejas (IDEAS • ESTRUCTURA • IMPACTO).
- **Estilo Visual:** Ejecutivo, neutral, sobrio y elegante. Sin retórica política, sin gráficos infantiles y sin iconos estilo años 80.

---

### 2. 🎨 SISTEMA DE DISEÑO (REMBRANDT & CARAVAGGIO CHIAROSCURO)
- **Fondo Web:** Negro absoluto sólido (`#000000`) en toda la plataforma. Sin patrones ni texturas.
- **Paleta Oficial:**
  - **Fondo General y Secciones:** `#000000` (Negro Absoluto).
  - **Tarjetas:** `#0D0D0D` con marco fino dorado `rgba(188, 167, 114, 0.25)`.
  - **Detalles Dorados:** `#BCA772` (Dorado Cálido / Oro Envejecido).
  - **Texto Principal:** `#FFFFFF` y `#E0E0E0`.
- **Tipografía Móvil:** Títulos de tarjetas en celulares fijados en **`15px !important` (Negrita)**.

---

### 3. 📐 ESTRUCTURA DE LA PORTADA (6 SECCIONES EN CARRUSEL HORIZONTAL)
La portada se compone de filas con deslizamiento horizontal en carrusel (flechas doradas `❮` `❯`):

1. **ÚLTIMAS ENTRADAS DEL BLOG:** Carrusel de artículos con botón *"VER TODAS ➔"*.
2. **SERVICIOS DE ESTRUCTURACIÓN:** *Destilado de Ideas*, *Forja de Proyectos*, *Marketing 360°*, *Impulso MYPE*.
3. **PROYECTOS PROPIOS & INVESTIGACIÓN:** *Modelo MFEIR*, *El Juego del Emprendedor*.
4. **LIBROS:** Sección dedicada a publicaciones y obras de Daniel Simons.
5. **QUIÉN SOY (DANIEL SIMONS):** Tarjeta ejecutiva de presentación.
6. **TRABAJOS Y PROPUESTAS:** *Propuesta Urbanizaciones*, *Transparencia Inteligente*, *Estrategia Electoral 2026*, *Guía Sobreviviendo a la Tesis*, *Evaluar Desarrollo de Marca*.

---

### 4. 📷 REGLAS DE IMÁGENES Y MÉTODO DE VINCULACIÓN CDN
- **Formato:** Miniaturas **100% cuadradas (1:1)** que ocupan toda la cabecera de la tarjeta (`object-fit: cover`).
- **MÉTODO ESTÁNDAR INMUTABLE PARA PÁGINAS:**
  Para cualquier cambio o adición a las páginas de *Servicios*, *Proyectos Propios*, *Trabajos y Propuestas* o cualquier página que se muestre en la portada principal, **SE UTILIZARÁ EXCLUSIVAMENTE EL MÉTODO DE VINCULACIÓN DIRECTA DESDE EL CDN DE BLOGGER (`blogger.googleusercontent.com`)**.
  - *Procedimiento:* Se toma la URL oficial de la imagen subida en esa página de Blogger y se vincula a la miniatura 1:1 de la tarjeta.
- **Prohibiciones Estrictas:**
  - 🚫 Prohibido iconos vectoriales o emojis antiguos.
  - 🚫 Prohibido fotos genéricas de stock.
  - 🚫 Prohibido colocar portadas de libros en filas de servicios.

---

### 5. ✍️ METODOLOGÍA OFICIAL DE PUBLICACIÓN DE ARTÍCULOS
La metodología perfecta a seguir de ahora en adelante será:

- **Redacción:** Yo redacto el artículo completo (con estructura AIDA) y te entrego el texto junto con la imagen recomendada.
- **Publicación por tu parte:** Tú entras a Blogger, pegas el texto, subes la foto con el botón *"Insertar imagen"* de Blogger y le das a Publicar.
- **Miniatura en Portada:** Me pasas el enlace del artículo publicado y yo coloco la miniatura cuadrada (1:1) oficial en el carrusel de la portada.

*Así se eliminan los cuadros plomos en tu editor, se evita cualquier imagen duplicada y mantienes el control 100% limpio sobre tus publicaciones.*

---

### 6. ⚡ LÍMITES TÉCNICOS Y PESO EN BLOGGER
- **Peso Máximo del XML:** El archivo XML del tema NUNCA DEBE SUPERAR LOS 200 KB (Rango ideal: 160 KB – 190 KB).
- **CLÁUSULA DE CONTROL ABSOLUTO:** Sólo modificarás LO QUE EXPRESAMENTE SE TE INSTRUYA; TODO LO DEMÁS QUEDARÁ EXACTAMENTE IGUAL.
"""

with open(md_path, "w", encoding="utf-8") as f:
    f.write(content_md)

# Save DOCX file
doc = docx.Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p_title = doc.add_paragraph()
run_title = p_title.add_run("PROMPT MASTER & REGLAS INMUTABLES DE LA WEB")
run_title.font.name = "Montserrat"
run_title.font.size = Pt(20)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(188, 167, 114)

p_sub = doc.add_paragraph()
run_sub = p_sub.add_run("Portal Oficial: www.danielsimons.xyz | Daniel Simons")
run_sub.font.name = "Montserrat"
run_sub.font.size = Pt(11)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph("="*60)

def add_sec(title, text_lines):
    h = doc.add_heading(title, level=2)
    for r in h.runs:
        r.font.name = "Montserrat"
        r.font.color.rgb = RGBColor(188, 167, 114)
    for line in text_lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Montserrat"
        r.font.size = Pt(10.5)

add_sec("1. POSICIONAMIENTO DE MARCA", [
    "• Marca: Daniel Simons — Estructurador de Ideas Complejas (IDEAS • ESTRUCTURA • IMPACTO).",
    "• Estilo Visual: Ejecutivo, neutral, sobrio y elegante. Sin retórica política, sin gráficos infantiles y sin iconos estilo años 80."
])

add_sec("2. SISTEMA DE DISEÑO (REMBRANDT & CARAVAGGIO CHIAROSCURO)", [
    "• Fondo Web: Negro absoluto sólido (#000000) en toda la plataforma. Sin patrones ni texturas.",
    "• Paleta Oficial: Fondo General (#000000), Tarjetas (#0D0D0D) con marco dorado (rgba(188, 167, 114, 0.25)), Detalles Dorados (#BCA772), Texto (#FFFFFF / #E0E0E0).",
    "• Tipografía Móvil: Títulos de tarjetas en celulares fijados en 15px !important (Negrita)."
])

add_sec("3. ESTRUCTURA DE LA PORTADA (CARRUSEL HORIZONTAL)", [
    "1. ÚLTIMAS ENTRADAS DEL BLOG: Carrusel de artículos con botón 'VER TODAS ➔'.",
    "2. SERVICIOS DE ESTRUCTURACIÓN: Destilado de Ideas, Forja de Proyectos, Marketing 360°, Impulso MYPE.",
    "3. PROYECTOS PROPIOS & INVESTIGACIÓN: Modelo MFEIR, El Juego del Emprendedor.",
    "4. LIBROS: Sección oficial de libros y publicaciones.",
    "5. QUIÉN SOY (DANIEL SIMONS): Tarjeta ejecutiva de presentación.",
    "6. TRABAJOS Y PROPUESTAS: Urbanizaciones, Transparencia Inteligente, Estrategia Electoral 2026, Guía Tesis, Evaluar Marca."
])

add_sec("4. REGLAS DE IMÁGENES Y MÉTODO DE VINCULACIÓN CDN", [
    "• Formato: Miniaturas 100% cuadradas (1:1) llenando la tarjeta (object-fit: cover).",
    "• MÉTODO ESTÁNDAR INMUTABLE: Para cualquier modificación o adición en las páginas de Servicios, Proyectos o Trabajos que vayan en la portada principal, SE UTILIZARÁ EXCLUSIVAMENTE EL MÉTODO DE VINCULACIÓN DIRECTA DESDE EL CDN DE BLOGGER (blogger.googleusercontent.com).",
    "• Prohibiciones Estrictas: Prohibido iconos noventeros, emojis antiguos, fotos genéricas de stock o portadas de libros en filas de servicios."
])

add_sec("5. METODOLOGÍA OFICIAL DE PUBLICACIÓN DE ARTÍCULOS", [
    "La metodología perfecta a seguir de ahora en adelante será:",
    "• Redacción: Yo redacto el artículo completo (con estructura AIDA) y te entrego el texto junto con la imagen recomendada.",
    "• Publicación por tu parte: Tú entras a Blogger, pegas el texto, subes la foto con el botón 'Insertar imagen' de Blogger y le das a Publicar.",
    "• Miniatura en Portada: Me pasas el enlace del artículo publicado y yo coloco la miniatura cuadrada (1:1) oficial en el carrusel de la portada.",
    "Así se eliminan los cuadros plomos en tu editor, se evita cualquier imagen duplicada y mantienes el control 100% limpio sobre tus publicaciones."
])

add_sec("6. LÍMITES TÉCNICOS Y PESO EN BLOGGER", [
    "• Peso Máximo del XML: NUNCA DEBE SUPERAR LOS 200 KB (Rango ideal: 160 KB – 190 KB).",
    "• CLÁUSULA DE CONTROL ABSOLUTO: Sólo modificarás LO QUE EXPRESAMENTE SE TE INSTRUYA; TODO LO DEMÁS QUEDARÁ EXACTAMENTE IGUAL."
])

doc.save(docx_path)
print("SUCCESS: Updated Master MD and DOCX files with Article Workflow!")
