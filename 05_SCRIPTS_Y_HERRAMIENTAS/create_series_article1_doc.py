import os
import docx
from docx.shared import Pt, RGBColor, Inches

root_dir = r"c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz"
art_dir = os.path.join(root_dir, "02_ARTICULOS_Y_PUBLICACIONES")

md_path = os.path.join(art_dir, "02.03_Articulo_Serie_1_Resolucion_Ministerial_245.md")
docx_path = os.path.join(art_dir, "02.03_Articulo_Serie_1_Resolucion_Ministerial_245.docx")

content_md = """# DE LA NORMA A LA CALLE: QUÉ DECRETA REALMENTE LA RESOLUCIÓN MINISTERIAL 245 Y CÓMO FUNCIONA EL NUEVO DÓLAR EN BOLIVIA
### Serie Política Monetaria y Sociedad | Parte 1
### Por Daniel Simons | Estructurador de Ideas Complejas

---

El 26 de junio de 2026 marca un punto de inflexión definitivo en la historia económica reciente de Bolivia. Tras más de una década bajo un tipo de cambio fijo e inamovible (6.96 Bs.), el Ministerio de Economía y Finanzas Públicas oficializó la **Resolución Ministerial N° 245**.

Aunque gran parte del debate público se concentra en el valor numérico diario del dólar en pizarras y redes sociales, muy pocos ciudadanos y empresarios han analizado con detenimiento el texto exacto de la norma. 

Comprender la arquitectura legal de este cambio no es un ejercicio académico; es la base indispensable para tomar decisiones financieras correctas.

---

### 📌 1. ATENCIÓN: El fin del mito del dólar fijo

Durante años, la población asumió que el valor de la moneda nacional era una constante inalterable. Sin embargo, la Resolución Ministerial N° 245 reconoce formalmente lo que la realidad económica venía mostrando: la caída estructural de los ingresos por exportación de gas y la reducción de las Reservas Internacionales Netas (RIN) volvieron insostenible la subvención estatal al tipo de cambio.

En su parte resolutiva, la norma marca un viraje de 180 grados mediante dos disposiciones fundamentales que todo ciudadano y emprendedor debe entender al detalle.

---

### 💡 2. INTERÉS: Los dos pilares de la Resolución Ministerial 245

#### ⚖️ Punto Primero: El establecimiento del Régimen Cambiario Flexible
El documento oficial establece formalmente el tránsito hacia un **Régimen Cambiario Flexible**. En términos sencillos, el Estado deja de fijar un precio artificial para el dólar. 

El objetivo normativo declarado es restablecer la competitividad externa, equilibrar la balanza de pagos y detener la sangría de divisas. En la práctica, esto significa que el valor del boliviano frente a monedas extranjeras pasa a ser una variable en movimiento constante.

#### 🏛️ Punto Segundo: La encomienda operativa al Banco Central de Bolivia (BCB)
La resolución instruye al Banco Central de Bolivia la ejecución diaria del tipo de cambio con una regla clara: **el reconocimiento de la oferta y la demanda real de divisas en el sistema financiero**.

Esto implica que el valor diario oficial ya no nace de un decreto estático, sino del promedio de las transacciones reales de compra y venta que realizan los bancos y entidades financieras. Si la demanda de dólares supera a la oferta disponible en el sistema bancario, la cotización tiende a ajustarse al alza; si ingresan divisas por exportaciones o créditos, la cotización busca estabilizarse.

---

### 🎯 3. DESEO: Las implicaciones inmediatas para tu bolsillo y tu negocio

Entender este mecanismo elimina la ilusión de que el dólar "volverá automáticamente a 6.96". Al pasar a un régimen flexible diario:

1. **Los costos de reposición cambian constantemente:** La mercadería e insumos importados se cotizan según el flujo diario del mercado financiero.
2. **El cálculo de márgenes debe ser dinámico:** Las empresas no pueden seguir fijando precios con costos históricos de compra.
3. **La certidumbre requiere análisis diario:** No basta con saber qué dice el periódico; es necesario comprender si el sistema financiero cuenta con la liquidez real para liquidar esas operaciones.

---

### 🚀 4. ACCIÓN: El siguiente paso en la discusión

La Resolución Ministerial N° 245 ha cambiado las reglas del juego. Sin embargo, haber pasado de un tipo de cambio fijo a una flotación libre de forma abrupta plantea un profundo debate técnico y social sobre su impacto real en la población.

¿Era la flotación libre inmediata la única opción, o existían mecanismos de ajuste gradual (como una flotación sucia administrada) para proteger el poder adquisitivo de las familias y el tejido de las MYPEs?

En la siguiente entrega de esta serie analizaremos a fondo el impacto social de este choque monetario y las alternativas técnicas de amortiguación.

👉 **[LEER LA PARTE 2: El dilema de la flotación libre vs. el ajuste gradual en la economía boliviana]**
"""

with open(md_path, "w", encoding="utf-8") as f:
    f.write(content_md)

# Create DOCX
doc = docx.Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p_title = doc.add_paragraph()
run_title = p_title.add_run("DE LA NORMA A LA CALLE: QUÉ DECRETA REALMENTE LA RESOLUCIÓN MINISTERIAL 245 Y CÓMO FUNCIONA EL NUEVO DÓLAR EN BOLIVIA")
run_title.font.name = "Montserrat"
run_title.font.size = Pt(15)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(188, 167, 114)

p_sub = doc.add_paragraph()
run_sub = p_sub.add_run("Serie Política Monetaria y Sociedad | Parte 1\nPor Daniel Simons | Estructurador de Ideas Complejas")
run_sub.font.name = "Montserrat"
run_sub.font.size = Pt(10.5)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph("="*60)

for line in content_md.split("\n"):
    if line.startswith("### "):
        h = doc.add_heading(line.replace("### ", ""), level=3)
        for r in h.runs:
            r.font.name = "Montserrat"
            r.font.color.rgb = RGBColor(188, 167, 114)
    elif line.startswith("#### "):
        h = doc.add_heading(line.replace("#### ", ""), level=4)
        for r in h.runs:
            r.font.name = "Montserrat"
            r.font.color.rgb = RGBColor(220, 220, 220)
    elif line.startswith("---") or line.startswith("# "):
        continue
    elif line.strip():
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Montserrat"
        r.font.size = Pt(10.5)

doc.save(docx_path)
print("SUCCESS: Series Article 1 docx and md created successfully!")
