import os
import docx
from docx.shared import Pt, RGBColor, Inches

root_dir = r"c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz"
master_md_path = os.path.join(root_dir, "00_DOCUMENTOS_ESTRATÉGICOS", "00.05_Prompt_Master_Diseno_y_Reglas_Web_DanielSimons.md")
master_docx_path = os.path.join(root_dir, "00_DOCUMENTOS_ESTRATÉGICOS", "00.05_Prompt_Master_Diseno_y_Reglas_Web_DanielSimons.docx")

content_md = """# 📜 PROMPT MASTER & REGLAS INMUTABLES DE LA WEB
### Portal Oficial: [www.danielsimons.xyz](https://www.danielsimons.xyz/)

---

### 1. 🎯 POSICIONAMIENTO DE MARCA
- **Marca:** Daniel Simons — Estructurador de Ideas Complejas (IDEAS • ESTRUCTURA • IMPACTO).
- **Estilo Visual:** Ejecutivo, neutral, sobrio y elegante. Sin retórica política, sin gráficos infantiles y sin iconos estilo años 80.

---

### 2. ⚡ AUTORIZACIÓN TOTAL DE EJECUCIÓN DE SCRIPTS Y PROTOCOLO "DALE"
- **EJECUCIÓN DIRECTA SIN PEDIR PERMISOS NUNCA:** Prohibido interrumpir con solicitudes emergentes o pedir permisos para ejecutar comandos de consola o scripts de Python. El asistente tiene **autorización absoluta y total** para ejecutar Python, comandos internos y actualizar archivos de forma fluida e ininterrumpida.
- **PROTOCOLO DE CONFIRMACIÓN CON "DALE":** Para avanzar a la redacción, modificación o publicación de contenido estratégico, el asistente presentará la propuesta e ideas clave y **esperará siempre que Daniel le diga "DALE" o le dé su visto bueno explícito** antes de proceder.

---

### 3. ✍️ GUÍA INMUTABLE DE ESTILO EDITORIAL DANIEL SIMONS
- **PÁRRAFOS CORTOS Y REDACCIÓN SISTÉMICA:**
  * **Párrafos Cortos:** Máximo 2 a 4 líneas por bloque. Cero bloques masivos de texto. Escaneo visual ultra-rápido en móviles y PCs.
  * **Redacción Sistémica:** Articular las ideas como un sistema interconectado de causa y efecto (Norma → Cambio Monetario → Liquidez Bancaria → Costo de Reposición → Ajuste Empresarial/Social).
- **CERO VERBORREA Y CERO HISTORIA ABURRIDA:** Prohibido prólogos académicos, introducciones históricas largas o discursos patronizantes. Entrada directa y quirúrgica desde la primera línea.
- **PROHIBICIÓN ABSOLUTA DE ICONOS, EMOJIS Y LÍNEAS SEPARADORAS (`---`):** NUNCA JAMÁS colocar emojis ni iconos en títulos o texto. NUNCA JAMÁS colocar rayas separadoras (`---` / `<hr/>`) dentro del cuerpo de los artículos.
- **TONO DE ALTA AUTORIDAD Y CERO CRÍTICA PATRONIZANTE:** Trato de igual a igual con el lector. Respeto absoluto al emprendedor y ciudadano. Enfoque en dar criterios de acción, no juzgar.
- **CONVERSIÓN UTILITARIA ORGÁNICA:** Cada escrito conduce de forma natural hacia el siguiente artículo o servicio de estructuración (*WhatsApp / Impulso MYPE / Destilado de Ideas*).

---

### 4. 📷 CARPETA DEDICADA DE IMÁGENES PARA PÁGINAS ESTÁTICAS Y PROTOCOLO 16:9
- **CARPETA DEDICADA `09_IMAGENES_PAGINAS_ESTATICAS/`:** Todas las imágenes de previsualización para páginas estáticas o paneles (*Indicadores Económicos, Servicios, Quiénes Somos, etc.*) se generarán y guardarán exclusivamente dentro de la carpeta dedicada `09_IMAGENES_PAGINAS_ESTATICAS/`.
- **FORMATO OFICIAL HORIZONTAL (16:9):** Se generarán en formato **Horizontal (16:9)** `.webp` con la estética fotorrealista macro (Alt2 + Alt5 Equipetrol).
- **PROTOCOLO DE SUBIDA MANUAL EN PÁGINAS ESTÁTICAS (`/p/*.html`):**
  * El asistente genera la imagen en `09_IMAGENES_PAGINAS_ESTATICAS/` y entrega la ruta directa a Daniel.
  * **Daniel sube la imagen manualmente al final de la página dentro del editor de Blogger** para que Blogger la registre nativamente como la miniatura oficial de previsualización en WhatsApp y redes sociales.
  * **Daniel comparte el enlace CDN oficial (`blogger.googleusercontent.com`)** para confirmación y vinculación.

---

### 5. 📐 ESTRUCTURA DE LA PORTADA (6 SECCIONES EN CARRUSEL HORIZONTAL)
La portada se compone de filas con deslizamiento horizontal en carrusel (flechas doradas `❮` `❯`):

1. **ÚLTIMAS ENTRADAS DEL BLOG:** Carrusel de artículos con botón *"Ver todas las entradas ➔"*.
2. **SERVICIOS DE ESTRUCTURACIÓN:** *Destilado de Ideas*, *FORJA*, *Marketing 360°*, *Impulso MYPE*.
3. **PROYECTOS PROPIOS & INVESTIGACIÓN:** *Modelo MFEIR*, *El Juego del Emprendedor*.
4. **LIBROS:** Sección dedicada a publicaciones y obras de Daniel Simons.
5. **QUIÉN SOY (DANIEL SIMONS):** Tarjeta ejecutiva de presentación.
6. **TRABAJOS Y PROPUESTAS:** *Propuesta Urbanizaciones*, *Transparencia Inteligente*, *Guía Sobreviviendo a la Tesis*, *Evaluar Desarrollo de Marca*.

---

### 6. 📁 ESTRUCTURA Y NÚCLEO EDITORIAL DE LIBROS: `07_JUEGO INFINITO/`
La obra completa de Daniel Simons se organiza en dos ejes principales:

- **`07_JUEGO INFINITO/ORIGINALES/`**: Contiene todo el archivo histórico y obras fuente originales clasificadas en:
  * `01_LIBROS/`
  * `02_MODELO_MEDS/` (Modelo de Estabilidad Dinámica Sistémica)
  * `03_MODELO_MFEIR/` (Modelo Computacional Basado en Agentes)
  * `04_INDIVIDUALISMO_RELACIONAL/` (Fundamento Filosófico Cero)
  * `05_RESCATADOS_Y_EVOLUTIVOS/`
  * `06_AGENTES_Y_PROMPTS/` (Inteligencia Dialéctica)

- **`07_JUEGO INFINITO/2026/`**: Espacio de trabajo oficial para las **Nuevas Versiones Avanzadas 2026** divididas en:
  * `01_LIBROS_2026/`
  * `02_MODELO_MEDS_2026/`
  * `03_MODELO_MFEIR_2026/`
  * `04_INDIVIDUALISMO_RELACIONAL_2026/`
  * `05_PRODUCTOS_Y_SERVICIOS_2026/`
  * `06_AGENTES_Y_PROMPTS_2026/`

---

### 7. ⚡ LÍMITES TÉCNICOS Y PESO EN BLOGGER
- **Peso Máximo del XML:** El archivo XML del tema NUNCA DEBE SUPERAR LOS 200 KB (Rango ideal: 160 KB – 190 KB).
- **CLÁUSULA DE CONTROL ABSOLUTO:** Sólo modificarás LO QUE EXPRESAMENTE SE TE INSTRUYA; TODO LO DEMÁS QUEDARÁ EXACTAMENTE IGUAL.
"""

with open(master_md_path, "w", encoding="utf-8") as f:
    f.write(content_md)

# Save Master DOCX
doc = docx.Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p_title = doc.add_paragraph()
run_title = p_title.add_run("PROMPT MASTER & REGLAS INMUTABLES DE LA WEB")
run_title.font.name = "Montserrat"
run_title.font.size = Pt(20)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(188, 167, 114)

p_sub = doc.add_paragraph()
run_sub = p_sub.add_run("Portal Oficial: www.danielsimons.xyz | Daniel Simons")
run_sub.font.name = "Montserrat"
run_sub.font.size = Pt(11)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph("="*60)

def add_sec(title, text_lines):
    h = doc.add_heading(title, level=2)
    for r in h.runs:
        r.font.name = "Montserrat"
        r.font.color.rgb = RGBColor(188, 167, 114)
    for line in text_lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Montserrat"
        r.font.size = Pt(10.5)

add_sec("1. POSICIONAMIENTO DE MARCA", [
    "• Marca: Daniel Simons — Estructurador de Ideas Complejas.",
    "• Estilo Visual: Ejecutivo, neutral, sobrio y elegante."
])

add_sec("2. AUTORIZACIÓN TOTAL DE EJECUCIÓN DE SCRIPTS Y PROTOCOLO 'DALE'", [
    "• EJECUCIÓN DIRECTA SIN PEDIR PERMISOS NUNCA.",
    "• PROTOCOLO DE CONFIRMACIÓN CON 'DALE'."
])

add_sec("3. GUÍA INMUTABLE DE ESTILO EDITORIAL DANIEL SIMONS", [
    "• PÁRRAFOS CORTOS: Máximo 2 a 4 líneas por bloque.",
    "• REDACCIÓN SISTÉMICA: Causa y efecto interconectados.",
    "• CERO EMOJIS Y CERO LÍNEAS SEPARADORAS (---)."
])

add_sec("4. CARPETA DEDICADA 09_IMAGENES_PAGINAS_ESTATICAS/ Y PROTOCOLO 16:9", [
    "• CARPETA DEDICADA: 09_IMAGENES_PAGINAS_ESTATICAS/.",
    "• SUBIDA MANUAL POR DANIEL EN BLOGGER."
])

add_sec("5. ESTRUCTURA DE LA PORTADA (CARRUSEL HORIZONTAL)", [
    "1. ÚLTIMAS ENTRADAS DEL BLOG",
    "2. SERVICIOS DE ESTRUCTURACIÓN",
    "3. PROYECTOS PROPIOS & INVESTIGACIÓN",
    "4. LIBROS",
    "5. QUIÉN SOY (DANIEL SIMONS)",
    "6. TRABAJOS Y PROPUESTAS"
])

add_sec("6. ESTRUCTURA Y NÚCLEO EDITORIAL: 07_JUEGO INFINITO/", [
    "• 07_JUEGO INFINITO/ORIGINALES/ (Materiales base y obras históricas)",
    "• 07_JUEGO INFINITO/2026/ (Nuevas versiones avanzadas 2026)"
])

add_sec("7. LÍMITES TÉCNICOS Y PESO EN BLOGGER", [
    "• Peso Máximo del XML: NUNCA DEBE SUPERAR LOS 200 KB."
])

doc.save(master_docx_path)
print("SUCCESS: Updated Master MD and DOCX with ORIGINALES & 2026 Structure!")
