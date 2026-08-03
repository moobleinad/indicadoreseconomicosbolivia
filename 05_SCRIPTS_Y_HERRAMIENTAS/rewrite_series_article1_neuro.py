import os
import docx
from docx.shared import Pt, RGBColor, Inches

root_dir = r"c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz"
art_dir = os.path.join(root_dir, "02_ARTICULOS_Y_PUBLICACIONES")

md_path = os.path.join(art_dir, "02.03_Articulo_Serie_1_Resolucion_Ministerial_245.md")
docx_path = os.path.join(art_dir, "02.03_Articulo_Serie_1_Resolucion_Ministerial_245.docx")

content_md = """# RÉGIMEN CAMBIARIO: CÓMO FUNCIONA AHORA EL TIPO DE CAMBIO DEL DÓLAR EN BOLIVIA
### Serie Política Monetaria y Sociedad | Parte 1
### Por Daniel Simons

---

El tipo de cambio fijo de 6.96 Bs. dejó de existir de forma oficial. Sin embargo, la mayoría de los negocios en Bolivia siguen intentando operar con reglas que ya no aplican. 

La **Resolución Ministerial N° 245** cambió el tablero. No es un ajuste estético: es la transición formal a un **Régimen Cambiario Flexible**. Si compras insumos, importas mercadería o fijas precios, necesitas entender los dos mandatos técnicos que rigen la economía desde hoy.

---

### ⚡ 1. Punto Primero: El fin de la cotización congelada

El Estado eliminó la fijación artificial del precio del dólar. Al establecer el régimen flexible, el valor del boliviano frente a divisas extranjeras pasa a ser una variable en constante movimiento.

- **Impacto directo:** Desaparece el costo histórico. El valor real de reposición de tus productos cambia según la disponibilidad del mercado.
- **Riesgo inmediato:** Fijar precios de venta con márgenes estáticos destruye tu flujo de caja en menos de 30 días.

---

### ⚡ 2. Punto Segundo: La orden operativa al Banco Central (BCB)

La norma instruye al BCB determinar el valor diario del dólar reconociendo la **oferta y demanda real del sistema financiero**.

- **El mecanismo:** El tipo de cambio oficial diario no nace de un escritorio estatal, sino del promedio de transacciones que logran cerrar bancos y entidades financieras.
- **La realidad operativa:** Si la demanda bancaria de divisas supera la oferta disponible, la cotización se desplaza. Si no hay dólares líquidos en ventanilla, el costo de transacción real se traslada a la operación privada.

---

### 📊 REGIONAL & EMPRESARIAL: ¿Qué hacer hoy con tu dinero?

1. **Cálculo por Costo de Reposición:** Reemplaza el costo de compra pasado por la cotización de reposición proyectada.
2. **Protección de Liquidez:** Monitorea la disponibilidad bancaria efectiva, no solo la cifra impresa en pizarra.
3. **Reajuste de Contratos:** Incorpora cláusulas de ajuste dinámico en ventas a crédito o insumos importados.

---

### 🔗 SIGUIENTE PASO EN ESTA SERIE:

¿Era la flotación libre inmediata la única salida o un suicidio de liquidez para el ciudadano y la MYPE? En la **Parte 2** analizamos la diferencia entre la terapia de choque actual y el modelo de flotación sucia administrada.

👉 **[IR A LA PARTE 2: El dilema de la flotación libre vs. flotación sucia en Bolivia]**

👉 **[EVALUAR ESTRUCTURA DE MI NEGOCIO CON DANIEL SIMONS](https://www.danielsimons.xyz/p/impulso-mype-360.html)**
"""

with open(md_path, "w", encoding="utf-8") as f:
    f.write(content_md)

# Save DOCX
doc = docx.Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p_title = doc.add_paragraph()
run_title = p_title.add_run("RÉGIMEN CAMBIARIO: CÓMO FUNCIONA AHORA EL TIPO DE CAMBIO DEL DÓLAR EN BOLIVIA")
run_title.font.name = "Montserrat"
run_title.font.size = Pt(16)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(188, 167, 114)

p_sub = doc.add_paragraph()
run_sub = p_sub.add_run("Serie Política Monetaria y Sociedad | Parte 1\nPor Daniel Simons")
run_sub.font.name = "Montserrat"
run_sub.font.size = Pt(11)
run_sub.font.bold = True
run_sub.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph("="*60)

for line in content_md.split("\n"):
    if line.startswith("### "):
        h = doc.add_heading(line.replace("### ", ""), level=3)
        for r in h.runs:
            r.font.name = "Montserrat"
            r.font.color.rgb = RGBColor(188, 167, 114)
    elif line.startswith("---") or line.startswith("# "):
        continue
    elif line.strip():
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Montserrat"
        r.font.size = Pt(10.5)

doc.save(docx_path)
print("SUCCESS: Rewrote Series Article 1 with Neuro-Direct Gerencial Style!")
