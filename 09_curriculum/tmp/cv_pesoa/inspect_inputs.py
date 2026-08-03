from pathlib import Path
import json
import pdfplumber
import pypdfium2 as pdfium
from pypdf import PdfReader
from docx import Document

ROOT = Path(r"C:\Users\Usuario\Desktop\daniel simons\curriculum")
TMP = ROOT / "tmp" / "cv_pesoa"
FILES = {
    "source": ROOT / "Copia de 8. ADM EMP ECO RESORT Daniel Simons 31 03 2026.pdf",
    "reference": ROOT / "convocatoria 1" / "CV Daniel Simons - Jefe de Marketing.pdf",
}

summary = {}
for name, path in FILES.items():
    out = TMP / name
    out.mkdir(parents=True, exist_ok=True)
    doc = pdfium.PdfDocument(str(path))
    pages = []
    for i in range(len(doc)):
        page = doc[i]
        bitmap = page.render(scale=1.7)
        image = bitmap.to_pil()
        image.save(out / f"page-{i+1}.png")
        pages.append({"page": i + 1, "width": image.width, "height": image.height})
    with pdfplumber.open(path) as pdf:
        text = "\n\n--- PAGE BREAK ---\n\n".join((p.extract_text(x_tolerance=2, y_tolerance=2) or "") for p in pdf.pages)
        imgs = [
            {"page": i + 1, "images": [
                {k: img.get(k) for k in ("x0", "x1", "top", "bottom", "width", "height", "name")}
                for img in p.images
            ]}
            for i, p in enumerate(pdf.pages)
        ]
    (TMP / f"{name}.txt").write_text(text, encoding="utf-8")
    summary[name] = {"path": str(path), "pages": pages, "image_objects": imgs}

docx_path = ROOT / "convocatoria 1" / "CV Daniel Simons - Jefe de Marketing.docx"
d = Document(docx_path)
docx_data = {
    "paragraphs": [{"i": i, "style": p.style.name, "text": p.text} for i, p in enumerate(d.paragraphs)],
    "tables": [
        {
            "table": ti,
            "rows": [[cell.text for cell in row.cells] for row in table.rows],
        }
        for ti, table in enumerate(d.tables)
    ],
    "inline_shapes": [
        {"width": s.width, "height": s.height, "type": str(s.type)}
        for s in d.inline_shapes
    ],
}
(TMP / "summary.json").write_text(json.dumps(summary, indent=2, ensure_ascii=False), encoding="utf-8")
(TMP / "reference_docx.json").write_text(json.dumps(docx_data, indent=2, ensure_ascii=False), encoding="utf-8")
print(json.dumps(summary, indent=2, ensure_ascii=False))
print(f"DOCX paragraphs={len(d.paragraphs)} tables={len(d.tables)} inline_shapes={len(d.inline_shapes)}")
