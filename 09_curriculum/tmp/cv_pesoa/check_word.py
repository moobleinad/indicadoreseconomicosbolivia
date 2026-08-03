from pathlib import Path
from zipfile import ZipFile
from docx import Document

p = Path(r"C:\Users\Usuario\Desktop\daniel simons\curriculum\convocatoria pesoa\CV Daniel Simons - Pesoa.docx")
d = Document(p)
text = "\n".join(x.text for x in d.paragraphs)
with ZipFile(p) as z:
    names = z.namelist()
    xml = z.read("word/document.xml").decode("utf-8")

print("archivo_bytes", p.stat().st_size)
print("parrafos", len(d.paragraphs))
print("tablas", len(d.tables))
print("imagenes", len([x for x in names if x.startswith("word/media/")]))
print("saltos_pagina", xml.count('w:type="page"'))
print("secciones", len(d.sections))
print("contenido_clave", all(x in text for x in [
    "EXPERIENCIA", "EDUCACIÓN", "OTROS", "Ing. Héctor Cáceres",
    "Certificación GOOGLE", "Diseño gráfico y de video"
]))
