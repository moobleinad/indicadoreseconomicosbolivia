from pathlib import Path
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\Usuario\Desktop\daniel simons\curriculum")
OUT_DIR = ROOT / "convocatoria pesoa"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "CV Daniel Simons - Pesoa.docx"
PHOTO = ROOT / "tmp" / "cv_pesoa" / "foto_daniel_simons.png"

RED = "B20D2A"
GRAY = RGBColor(102, 102, 102)
BLACK = RGBColor(17, 17, 17)


def set_font(run, name="Arial", size=9.5, bold=False, italic=False, color=BLACK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def no_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)


def add_bottom_border(paragraph, color=RED, size=8):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_font(run, size=7, color=GRAY)


doc = Document()
sec = doc.sections[0]
sec.page_width = Mm(210)
sec.page_height = Mm(297)
sec.left_margin = Mm(18)
sec.right_margin = Mm(18)
sec.top_margin = Mm(19)
sec.bottom_margin = Mm(16)
sec.header_distance = Mm(8)
sec.footer_distance = Mm(8)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(9.5)
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.05

heading = doc.styles["Heading 1"]
heading.font.name = "Arial"
heading._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
heading.font.size = Pt(11.5)
heading.font.bold = True
heading.font.color.rgb = RGBColor(178, 13, 42)
heading.paragraph_format.space_before = Pt(7)
heading.paragraph_format.space_after = Pt(4)
heading.paragraph_format.keep_with_next = True

bullet_style = doc.styles["List Bullet"]
bullet_style.font.name = "Arial"
bullet_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
bullet_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
bullet_style.font.size = Pt(9.25)
bullet_style.paragraph_format.left_indent = Mm(6)
bullet_style.paragraph_format.first_line_indent = Mm(-3.5)
bullet_style.paragraph_format.space_after = Pt(1)
bullet_style.paragraph_format.line_spacing = 1.0

# Header and footer
hp = sec.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("DANIEL SIMONS  |  CURRÍCULUM VITAE")
set_font(hr, size=7.3, bold=True, color=RGBColor(128, 128, 128))

fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Daniel Simons   ·   +591 65050351   ·   Página ")
set_font(fr, size=7, color=GRAY)
add_page_field(fp)


def add_section(text):
    par = doc.add_paragraph(text, style="Heading 1")
    add_bottom_border(par)
    return par


def add_rich(lines, compact=True):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(3 if compact else 5)
    par.paragraph_format.line_spacing = 1.0
    for i, (text, bold) in enumerate(lines):
        if i:
            par.add_run().add_break()
        run = par.add_run(text)
        set_font(run, size=9.25, bold=bold)
    return par


def add_bullet(text):
    par = doc.add_paragraph(style="List Bullet")
    run = par.add_run(text)
    set_font(run, size=9.25)
    return par


# Page 1 title block, using a two-cell header table for an editable layout.
table = doc.add_table(rows=1, cols=2)
table.autofit = False
table.columns[0].width = Mm(145)
table.columns[1].width = Mm(27)
no_table_borders(table)
for cell in table.rows[0].cells:
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
left = table.cell(0, 0)
right = table.cell(0, 1)
left.width = Mm(145)
right.width = Mm(27)
lp = left.paragraphs[0]
lp.paragraph_format.space_after = Pt(1)
nr = lp.add_run("DANIEL SIMONS")
set_font(nr, size=25, bold=True)
cp = left.add_paragraph()
cp.paragraph_format.space_after = Pt(1)
for idx, text in enumerate([
    "danielbsimons@gmail.com   ·   +591 65050351",
    "B. Est. Argentina / C. José V. Solíz No 3070   ·   48 AÑOS",
]):
    if idx:
        cp.add_run().add_break()
    cr = cp.add_run(text)
    set_font(cr, size=9.3, color=GRAY)
rp = right.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rp.add_run().add_picture(str(PHOTO), width=Mm(21), height=Mm(23))

add_section("EXPERIENCIA")
add_rich([("Sinergia Solutions - Tecnología en Refrigeración (2023-2025)", True),
          ("Administrativo comercial", False), ("Santa Cruz - Bolivia", False)])
add_rich([("Administrador comercial (2021-2023)", True),
          ("Administrador de personal de ventas y marketing.", False),
          ("Ofertangas - Santa Cruz - Bolivia", False)])
add_rich([("Analista de planificación y proyectos (2018-2020)", True),
          ("Grupo SION Inmobiliaria Kintas s.r.l.", False), ("Santa Cruz, Bolivia", False)])
add_rich([("Coordinación Administrativa (2015-2018)", True),
          ("Departamento de planificación", False),
          ("Secretaría Municipal de Seguridad Ciudadana - Gobierno Autónomo Municipal de Santa Cruz", False),
          ("Santa Cruz, Bolivia", False)])
add_rich([("Anteriormente:", True)])
for item in [
    "Administrativo y auxiliar administrativo en Constructora Quebracho. (Manejo de personal)",
    "Asesor de micro negocios en la empresa Servidor S.A.",
    "Asistente técnico para la Escuela de Proyectos de la CIDOB.",
    "Diseño y administración de campañas publicitarias en redes sociales y Google.",
]:
    add_bullet(item)
add_rich([("Docente de Taller de ideas de negocio Instituto ILACE. Docente de Proyectos de emprendimientos "
           "productivos Instituto ILACE. Docente de Taller de modalidad de grado Instituto ILACE.", False)])
add_section("SKILLS")
for item in ["Autogestión y trabajo por metas.", "Dinámico y versátil.", "Metódico, planificado y organizado."]:
    add_bullet(item)
add_section("REFERENCIAS")
add_rich([("Ing. Héctor Cáceres", True), ("GERENTE GENERAL - SINERGIA SOLUTIONS", False), ("72606967", False)])
add_rich([("Ing. Gabriel Simons", True), ("Jefe de producción - ALICORP", False), ("72131765", False)])

doc.add_page_break()
add_section("EDUCACIÓN")
for title, detail in [
    ("Cursos de POSTGRADO en Formulación y Gestión de Proyectos (2013-2015)", "Escuela de Negocios UAGRMBS\nSanta Cruz, Bolivia"),
    ("Economía (Licenciatura)", "Universidad Autónoma Gabriel René Moreno\nSanta Cruz, Bolivia"),
    ("Curso de Gestión de Alcance, Tiempo y Costo (PMI - PMBOK) (2013)", "Escuela de Negocios UAGRMBS\nSanta Cruz, Bolivia"),
    ("Curso de Gestión de calidad, Recursos humanos y comunicaciones (PMI - PMBOK) (2014)", "Escuela de Negocios UAGRMBS\nSanta Cruz, Bolivia"),
    ("Curso de Gestión de riesgos, adquisiciones, interesados (PMI - PMBOK) (2014)", "Escuela de Negocios UAGRMBS\nSanta Cruz, Bolivia"),
    ("Curso de Fundamentos de la escritura en español (2016)", "Tecnológico de Monterrey - Coursera (On line)\nSanta Cruz, Bolivia"),
    ("Curso especializado de MARKETING GERENCIAL (Curso) 2021", "Universidad de Chile - Coursera (On line)\nSanta Cruz, Bolivia"),
]:
    parts = [(title, True)] + [(x, False) for x in detail.split("\n")]
    add_rich(parts)
add_rich([("Certificación GOOGLE, de Marketing DIGITAL al 100 % completado (Demostrable).", True)])
add_rich([("Certificación ADS Expert by Aleph Consulting al 60 % de avance (Demostrable).", True)])
add_section("SKILLS")
for item in ["Autogestión y trabajo por metas.", "Dinámico y versátil.", "Metódico, planificado y organizado."]:
    add_bullet(item)
add_section("REFERENCIAS")
add_rich([("Ing. Héctor Cáceres", True), ("GERENTE GENERAL - SINERGIA SOLUTIONS", False), ("72606967", False)])
add_rich([("Ing. Gabriel Simons", True), ("Jefe de producción - ALICORP", False), ("72131765", False)])

doc.add_page_break()
add_section("OTROS")
add_rich([("HABILIDADES:", True)])
for item in [
    "Gestión y planificación certificada.", "Buen manejo financiero y de proyectos.",
    "Buen manejo de herramientas publicitarias “ADS”.", "Operador de Microsoft Office.",
    "Excelente redacción.", "Análisis numérico y analítica.",
    "Buen manejo y rápido aprendizaje de softwares.", "Diseño gráfico y de video.",
]:
    add_rich([(item, False)])
add_rich([("IDIOMAS:", True)])
add_rich([("Inglés   Nivel Oral: Básico   Nivel Escrito: Básico   Nivel Lectura: Medio alto", False)])
add_section("REFERENCIAS")
for name, role, phone in [
    ("Ing. Carlos Simons Rocha", "Ex gerente de Constructora Quebracho", "Celular: 76020132"),
    ("Ing. Pablo Sesgua", "Ex Jefe de Planificación Inmobiliaria Kintas Grupo SION", "Celular: 72690937"),
    ("Lic. José Negrete", "Ex Secretario de Seguridad Ciudadana Gobierno Autónomo Municipal SCZ", "Celular: 77395905"),
    ("Lic. Jaime Vargas", "Ex Director de Cultura de la Casa del Cultura Santa Cruz", "Celular: 79878767"),
    ("Ing. Gabriel Simons Rocha", "Jefe de Producción ALICORP", "Celular: 72131765"),
    ("Ing. Héctor Cáceres", "GERENTE GENERAL SINERGIA SOLUTIONS", "Celular: 72606967"),
]:
    add_rich([(name, True), (role, False), (phone, False)])

doc.core_properties.title = "Currículum Vitae - Daniel Simons"
doc.core_properties.author = "Daniel Simons"
doc.save(OUT)
print(OUT)
