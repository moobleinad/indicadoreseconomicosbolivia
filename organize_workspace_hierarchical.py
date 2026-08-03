import os
import shutil
import docx
from docx.shared import Pt, RGBColor, Inches

root_dir = r"c:\Users\Usuario\Desktop\antigravity\blogger_danielsimons.xyz"

# 1. DEFINE TARGET DIRECTORY STRUCTURE WITH PLAIN ASCII / CLEAN NAMES
dirs_to_create = [
    "00_DOCUMENTOS_ESTRATEGICOS",
    "01_LINEA_GRAFICA_Y_ASSETS",
    "02_ARTICULOS_Y_PUBLICACIONES",
    "03_TEMAS_Y_PLANTILLAS_XML",
    "04_API_BLOGGER_Y_AUTENTICACION",
    "05_SCRIPTS_Y_HERRAMIENTAS",
    "06_ARCHIVOS_TEMPORALES_Y_CACHE"
]

for d in dirs_to_create:
    os.makedirs(os.path.join(root_dir, d), exist_ok=True)

# MAPPING OF FILES / DIRECTORIES TO THEIR NUMBERED CATEGORIES
move_map = {
    # 00_DOCUMENTOS_ESTRATEGICOS
    "0_Perfil_Contexto_DanielSimons.md": "00_DOCUMENTOS_ESTRATEGICOS/00.01_Perfil_Contexto_DanielSimons.md",
    "0_Modelo_de_Negocio_DanielSimons.md": "00_DOCUMENTOS_ESTRATEGICOS/00.02_Modelo_de_Negocio_DanielSimons.md",
    "0_Informe_Auditoria_DanielSimons.md": "00_DOCUMENTOS_ESTRATEGICOS/00.03_Informe_Auditoria_DanielSimons.md",
    "0_Informe_Auditoria_DanielSimons.docx": "00_DOCUMENTOS_ESTRATEGICOS/00.03_Informe_Auditoria_DanielSimons.docx",
    "0_Plan_Estrategico_Posicionamiento_y_Contenidos_DanielSimons.md": "00_DOCUMENTOS_ESTRATEGICOS/00.04_Plan_Estrategico_Posicionamiento_y_Contenidos_DanielSimons.md",
    "0_Plan_Estrategico_Posicionamiento_y_Contenidos_DanielSimons.docx": "00_DOCUMENTOS_ESTRATEGICOS/00.04_Plan_Estrategico_Posicionamiento_y_Contenidos_DanielSimons.docx",
    "0_Prompt_Master_Diseno_y_Reglas_Web_DanielSimons.md": "00_DOCUMENTOS_ESTRATEGICOS/00.05_Prompt_Master_Diseno_y_Reglas_Web_DanielSimons.md",
    "0_Prompt_Master_Diseno_y_Reglas_Web_DanielSimons.docx": "00_DOCUMENTOS_ESTRATEGICOS/00.05_Prompt_Master_Diseno_y_Reglas_Web_DanielSimons.docx",
    "0_Guia_Buenas_Practicas_Codigo_SEO_y_Rendimiento_DanielSimons.md": "00_DOCUMENTOS_ESTRATEGICOS/00.06_Guia_Buenas_Practicas_Codigo_SEO_y_Rendimiento_DanielSimons.md",
    "0_Guia_Integracion_Blogger_API.md": "00_DOCUMENTOS_ESTRATEGICOS/00.07_Guia_Integracion_Blogger_API.md",

    # 01_LINEA_GRAFICA_Y_ASSETS
    "linea grafica": "01_LINEA_GRAFICA_Y_ASSETS/01.01_linea_grafica",
    "foto_articulo1_bth_oficial.webp": "01_LINEA_GRAFICA_Y_ASSETS/01.02_foto_articulo1_bth_oficial.webp",
    "foto_articulo1_bth_cuadrada.jpg": "01_LINEA_GRAFICA_Y_ASSETS/01.03_foto_articulo1_bth_cuadrada.jpg",
    "maqueta_diseno.jpg": "01_LINEA_GRAFICA_Y_ASSETS/01.04_maqueta_diseno.jpg",
    "maqueta_inicio.html": "01_LINEA_GRAFICA_Y_ASSETS/01.05_maqueta_inicio.html",

    # 02_ARTICULOS_Y_PUBLICACIONES
    "0_Articulo_1_El_Drama_del_BTH_DanielSimons.md": "02_ARTICULOS_Y_PUBLICACIONES/02.01_Articulo_1_El_Drama_del_BTH_DanielSimons.md",
    "0_Articulo_1_El_Drama_del_BTH_DanielSimons.docx": "02_ARTICULOS_Y_PUBLICACIONES/02.01_Articulo_1_El_Drama_del_BTH_DanielSimons.docx",
    "ENTRADAS DESCARTADAS": "02_ARTICULOS_Y_PUBLICACIONES/02.02_ENTRADAS_DESCARTADAS",
    "0 INDIVIDUALISMO RELACIONAL": "02_ARTICULOS_Y_PUBLICACIONES/02.03_0_INDIVIDUALISMO_RELACIONAL",
    "JUEGO INFINITO": "02_ARTICULOS_Y_PUBLICACIONES/02.04_JUEGO_INFINITO",

    # 03_TEMAS_Y_PLANTILLAS_XML
    "copia de seguridad del tema web": "03_TEMAS_Y_PLANTILLAS_XML/03.01_copia_de_seguridad_del_tema_web",
    "tema_optimizado": "03_TEMAS_Y_PLANTILLAS_XML/03.02_tema_optimizado",

    # 04_API_BLOGGER_Y_AUTENTICACION
    "credencial blogger": "04_API_BLOGGER_Y_AUTENTICACION/04.01_credencial_blogger",
    "client_secret.json": "04_API_BLOGGER_Y_AUTENTICACION/04.02_client_secret.json",
    "token.json": "04_API_BLOGGER_Y_AUTENTICACION/04.03_token.json",
    "auth_url.txt": "04_API_BLOGGER_Y_AUTENTICACION/04.04_auth_url.txt",
}

# MOVE RECOGNIZED FILES AND DIRECTORIES SAFELY
for src_name, rel_dest in move_map.items():
    src_path = os.path.join(root_dir, src_name)
    dest_path = os.path.join(root_dir, rel_dest)
    if os.path.exists(src_path):
        try:
            if os.path.isdir(src_path):
                if not os.path.exists(dest_path):
                    shutil.copytree(src_path, dest_path)
                try:
                    shutil.rmtree(src_path)
                except Exception as e:
                    print(f"Directory kept or locked: {src_name}")
            else:
                shutil.move(src_path, dest_path)
            print(f"Processed: {src_name} -> {rel_dest}")
        except Exception as err:
            print(f"Skip {src_name}: {err}")

# MOVE ALL PYTHON SCRIPTS TO 05_SCRIPTS_Y_HERRAMIENTAS
for f in os.listdir(root_dir):
    full_path = os.path.join(root_dir, f)
    if os.path.isfile(full_path) and f.endswith(".py") and f != "organize_workspace_hierarchical.py":
        dest_path = os.path.join(root_dir, "05_SCRIPTS_Y_HERRAMIENTAS", f)
        try:
            shutil.move(full_path, dest_path)
            print(f"Moved Script: {f} -> 05_SCRIPTS_Y_HERRAMIENTAS/{f}")
        except Exception:
            pass
    elif os.path.isfile(full_path) and (f.endswith(".jpg") or f.endswith(".json")):
        dest_path = os.path.join(root_dir, "06_ARCHIVOS_TEMPORALES_Y_CACHE", f)
        try:
            shutil.move(full_path, dest_path)
            print(f"Moved Temp/Cache File: {f} -> 06_ARCHIVOS_TEMPORALES_Y_CACHE/{f}")
        except Exception:
            pass

print("SUCCESS: Workspace Re-organization Complete!")

# UPDATE MASTER PROMPT MD & DOCX IN THEIR NEW PATH
master_md_path = os.path.join(root_dir, "00_DOCUMENTOS_ESTRATEGICOS/00.05_Prompt_Master_Diseno_y_Reglas_Web_DanielSimons.md")
master_docx_path = os.path.join(root_dir, "00_DOCUMENTOS_ESTRATEGICOS/00.05_Prompt_Master_Diseno_y_Reglas_Web_DanielSimons.docx")

content_md = """# 📜 PROMPT MASTER & REGLAS INMUTABLES DE LA WEB
### Portal Oficial: [www.danielsimons.xyz](https://www.danielsimons.xyz/)

---

### 1. 🎯 POSICIONAMIENTO DE MARCA
- **Marca:** Daniel Simons — Estructurador de Ideas Complejas (IDEAS • ESTRUCTURA • IMPACTO).
- **Estilo Visual:** Ejecutivo, neutral, sobrio y elegante. Sin retórica política, sin gráficos infantiles y sin iconos estilo años 80.

---

### 2. 🎨 SISTEMA DE DISEÑO (REMBRANDT & CARAVAGGIO CHIAROSCURO)
- **Fondo Web:** Negro absoluto sólido (`#000000`) en toda la plataforma. Sin patrones ni texturas.
- **Paleta Oficial:**
  - **Fondo General y Secciones:** `#000000` (Negro Absoluto).
  - **Tarjetas:** `#0D0D0D` con marco fino dorado `rgba(188, 167, 114, 0.25)`.
  - **Detalles Dorados:** `#BCA772` (Dorado Cálido / Oro Envejecido).
  - **Texto Principal:** `#FFFFFF` y `#E0E0E0`.
- **Tipografía Móvil:** Títulos de tarjetas en celulares fijados en **`15px !important` (Negrita)**.

---

### 3. 📐 ESTRUCTURA DE LA PORTADA (6 SECCIONES EN CARRUSEL HORIZONTAL)
La portada se compone de filas con deslizamiento horizontal en carrusel (flechas doradas `❮` `❯`):

1. **ÚLTIMAS ENTRADAS DEL BLOG:** Carrusel de artículos con botón *"VER TODAS ➔"*.
2. **SERVICIOS DE ESTRUCTURACIÓN:** *Destilado de Ideas*, *Forja de Proyectos*, *Marketing 360°*, *Impulso MYPE*.
3. **PROYECTOS PROPIOS & INVESTIGACIÓN:** *Modelo MFEIR*, *El Juego del Emprendedor*.
4. **LIBROS:** Sección dedicada a publicaciones y obras de Daniel Simons.
5. **QUIÉN SOY (DANIEL SIMONS):** Tarjeta ejecutiva de presentation.
6. **TRABAJOS Y PROPUESTAS:** *Propuesta Urbanizaciones*, *Transparencia Inteligente*, *Estrategia Electoral 2026*, *Guía Sobreviviendo a la Tesis*, *Evaluar Desarrollo de Marca*.

---

### 4. 📷 REGLAS DE IMÁGENES Y MÉTODO DE VINCULACIÓN CDN
- **Formato:** Miniaturas **100% cuadradas (1:1)** que ocupan toda la cabecera de la tarjeta (`object-fit: cover`).
- **MÉTODO ESTÁNDAR INMUTABLE PARA PÁGINAS:**
  Para cualquier cambio o adición a las páginas de *Servicios*, *Proyectos Propios*, *Trabajos y Propuestas* o cualquier página que se muestre en la portada principal, **SE UTILIZARÁ EXCLUSIVAMENTE EL MÉTODO DE VINCULACIÓN DIRECTA DESDE EL CDN DE BLOGGER (`blogger.googleusercontent.com`)**.
  - *Procedimiento:* Se toma la URL oficial de la imagen subida en esa página de Blogger y se vincula a la miniatura 1:1 de la tarjeta.
- **Prohibiciones Estrictas:**
  - 🚫 Prohibido iconos vectoriales o emojis antiguos.
  - 🚫 Prohibido fotos genéricas de stock.
  - 🚫 Prohibido colocar portadas de libros en filas de servicios.

---

### 5. ✍️ METODOLOGÍA OFICIAL DE PUBLICACIÓN DE ARTÍCULOS
La metodología perfecta a seguir de ahora en adelante será:

- **Redacción:** Yo redacto el artículo completo (con estructura AIDA) y te entrego el texto junto con la imagen recomendada.
- **Publicación por tu parte:** Tú entras a Blogger, pegas el texto, subes la foto con el botón *"Insertar imagen"* de Blogger y le das a Publicar.
- **Miniatura en Portada:** Me pasas el enlace del artículo publicado y yo coloco la miniatura cuadrada (1:1) oficial en el carrusel de la portada.

---

### 6. 📁 SISTEMA NUMÉRICO Y NORMAS DE ORGANIZACIÓN DE ARCHIVOS
- **Jerarquía Numérica:** Carpetas y archivos organizados de forma secuencial cronológica y por categorías con un número por delante (`00_...`, `01_...`, `02_...`).
- **Subdivisión Decimal:** Los documentos internos se numeran decimalmente como un índice (`00.01_...`, `00.02_...`, `01.01_...`) para mantener el orden cronológico y jerárquico.
- **Estructura Oficial del Proyecto:**
  - `00_DOCUMENTOS_ESTRATEGICOS/` (Modelos, Planes, Prompts y Auditorías).
  - `01_LINEA_GRAFICA_Y_ASSETS/` (Guías gráficas, maquetas, imágenes WebP).
  - `02_ARTICULOS_Y_PUBLICACIONES/` (Redacciones de artículos, borradores, libros).
  - `03_TEMAS_Y_PLANTILLAS_XML/` (Copias de seguridad y versiones `v17`, `v18`).
  - `04_API_BLOGGER_Y_AUTENTICACION/` (Credenciales y tokens de acceso).
  - `05_SCRIPTS_Y_HERRAMIENTAS/` (Scripts de automatización y parches en Python).
  - `06_ARCHIVOS_TEMPORALES_Y_CACHE/` (Cachés, descargas y temporales).

---

### 7. ⚡ LÍMITES TÉCNICOS Y PESO EN BLOGGER
- **Peso Máximo del XML:** El archivo XML del tema NUNCA DEBE SUPERAR LOS 200 KB (Rango ideal: 160 KB – 190 KB).
- **CLÁUSULA DE CONTROL ABSOLUTO:** Sólo modificarás LO QUE EXPRESAMENTE SE TE INSTRUYA; TODO LO DEMÁS QUEDARÁ EXACTAMENTE IGUAL.
"""

with open(master_md_path, "w", encoding="utf-8") as f:
    f.write(content_md)

# Save Master DOCX
doc = docx.Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

p_title = doc.add_paragraph()
run_title = p_title.add_run("PROMPT MASTER & REGLAS INMUTABLES DE LA WEB")
run_title.font.name = "Montserrat"
run_title.font.size = Pt(20)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(188, 167, 114)

p_sub = doc.add_paragraph()
run_sub = p_sub.add_run("Portal Oficial: www.danielsimons.xyz | Daniel Simons")
run_sub.font.name = "Montserrat"
run_sub.font.size = Pt(11)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph("="*60)

def add_sec(title, text_lines):
    h = doc.add_heading(title, level=2)
    for r in h.runs:
        r.font.name = "Montserrat"
        r.font.color.rgb = RGBColor(188, 167, 114)
    for line in text_lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Montserrat"
        r.font.size = Pt(10.5)

add_sec("1. POSICIONAMIENTO DE MARCA", [
    "• Marca: Daniel Simons — Estructurador de Ideas Complejas (IDEAS • ESTRUCTURA • IMPACTO).",
    "• Estilo Visual: Ejecutivo, neutral, sobrio y elegante. Sin retórica política, sin gráficos infantiles y sin iconos estilo años 80."
])

add_sec("2. SISTEMA DE DISEÑO (REMBRANDT & CARAVAGGIO CHIAROSCURO)", [
    "• Fondo Web: Negro absoluto sólido (#000000) en toda la plataforma. Sin patrones ni texturas.",
    "• Paleta Oficial: Fondo General (#000000), Tarjetas (#0D0D0D) con marco dorado (rgba(188, 167, 114, 0.25)), Detalles Dorados (#BCA772), Texto (#FFFFFF / #E0E0E0).",
    "• Tipografía Móvil: Títulos de tarjetas en celulares fijados en 15px !important (Negrita)."
])

add_sec("3. ESTRUCTURA DE LA PORTADA (CARRUSEL HORIZONTAL)", [
    "1. ÚLTIMAS ENTRADAS DEL BLOG: Carrusel de artículos con botón 'VER TODAS ➔'.",
    "2. SERVICIOS DE ESTRUCTURACIÓN: Destilado de Ideas, Forja de Proyectos, Marketing 360°, Impulso MYPE.",
    "3. PROYECTOS PROPIOS & INVESTIGACIÓN: Modelo MFEIR, El Juego del Emprendedor.",
    "4. LIBROS: Sección oficial de libros y publicaciones.",
    "5. QUIÉN SOY (DANIEL SIMONS): Tarjeta ejecutiva de presentación.",
    "6. TRABAJOS Y PROPUESTAS: Urbanizaciones, Transparencia Inteligente, Estrategia Electoral 2026, Guía Tesis, Evaluar Marca."
])

add_sec("4. REGLAS DE IMÁGENES Y MÉTODO DE VINCULACIÓN CDN", [
    "• Formato: Miniaturas 100% cuadradas (1:1) llenando la tarjeta (object-fit: cover).",
    "• MÉTODO ESTÁNDAR INMUTABLE: Para cualquier modificación o adición en las páginas de Servicios, Proyectos o Trabajos que vayan en la portada principal, SE UTILIZARÁ EXCLUSIVAMENTE EL MÉTODO DE VINCULACIÓN DIRECTA DESDE EL CDN DE BLOGGER (blogger.googleusercontent.com).",
    "• Prohibiciones Estrictas: Prohibido iconos noventeros, emojis antiguos, fotos genéricas de stock o portadas de libros en filas de servicios."
])

add_sec("5. METODOLOGÍA OFICIAL DE PUBLICACIÓN DE ARTÍCULOS", [
    "• Redacción: Yo redacto el artículo completo (con estructura AIDA) y te entrego el texto junto con la imagen recomendada.",
    "• Publicación por tu parte: Tú entras a Blogger, pegas el texto, subes la foto con el botón 'Insertar imagen' de Blogger y le das a Publicar.",
    "• Miniatura en Portada: Me pasas el enlace del artículo publicado y yo coloco la miniatura cuadrada (1:1) oficial en el carrusel de la portada."
])

add_sec("6. SISTEMA NUMÉRICO Y NORMAS DE ORGANIZACIÓN DE ARCHIVOS", [
    "• Jerarquía Numérica: Carpetas y archivos organizados de forma secuencial cronológica y por categorías con un número por delante (00_..., 01_..., 02_...).",
    "• Subdivisión Decimal: Los documentos internos se numeran decimalmente como un índice (00.01_..., 00.02_..., 01.01_...) para mantener el orden cronológico y jerárquico.",
    "• Estructura Oficial del Proyecto: 00_DOCUMENTOS_ESTRATEGICOS, 01_LINEA_GRAFICA_Y_ASSETS, 02_ARTICULOS_Y_PUBLICACIONES, 03_TEMAS_Y_PLANTILLAS_XML, 04_API_BLOGGER_Y_AUTENTICACION, 05_SCRIPTS_Y_HERRAMIENTAS, 06_ARCHIVOS_TEMPORALES_Y_CACHE."
])

add_sec("7. LÍMITES TÉCNICOS Y PESO EN BLOGGER", [
    "• Peso Máximo del XML: NUNCA DEBE SUPERAR LOS 200 KB (Rango ideal: 160 KB – 190 KB).",
    "• CLÁUSULA DE CONTROL ABSOLUTO: Sólo modificarás LO QUE EXPRESAMENTE SE TE INSTRUYA; TODO LO DEMÁS QUEDARÁ EXACTAMENTE IGUAL."
])

doc.save(master_docx_path)
print("SUCCESS: Updated Master MD and DOCX with File System Rule!")
